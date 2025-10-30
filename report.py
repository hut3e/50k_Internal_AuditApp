import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import io
import base64
from datetime import datetime
import numpy as np
import json
import sys
import traceback
import os

from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Thêm vào đầu file - Thay đổi sang fpdf2 thay vì fpdf
from fpdf import FPDF
import urllib.request
# Thêm thư viện để hỗ trợ Unicode
import pkg_resources

from database_helper import get_supabase_client

# Giả lập database_helper nếu không có
try:
    from database_helper import check_answer_correctness, get_all_questions, get_all_users, get_user_submissions, get_all_submissions
except ImportError:
    # Mock functions để tránh lỗi khi không có module
    def check_answer_correctness(user_ans, q):
        """Mock function - nên sử dụng hàm từ database_helper"""
        if not user_ans:
            return False
        
        q_type = q.get("type", "")
        
        # Câu hỏi tự luận (Essay): tính đúng nếu có nội dung
        if q_type == "Essay":
            return bool(user_ans) and isinstance(user_ans[0], str) and user_ans[0].strip() != ""
        
        # Câu hỏi Combobox: chọn một đáp án
        if q_type == "Combobox":
            if len(user_ans) == 1:
                answer_text = user_ans[0]
                answers = q.get("answers", [])
                correct = q.get("correct", [])
                if isinstance(answers, str):
                    try:
                        answers = json.loads(answers)
                    except:
                        answers = [answers]
                if isinstance(correct, str):
                    try:
                        correct = json.loads(correct)
                    except:
                        try:
                            correct = [int(x.strip()) for x in correct.split(",")]
                        except:
                            correct = []
                answer_index = answers.index(answer_text) + 1 if answer_text in answers else -1
                return answer_index in correct
            return False
        
        # Câu hỏi Checkbox: nhiều lựa chọn
        if q_type == "Checkbox":
            answers = q.get("answers", [])
            if isinstance(answers, str):
                try:
                    answers = json.loads(answers)
                except:
                    answers = [answers]
            correct = q.get("correct", [])
            if isinstance(correct, str):
                try:
                    correct = json.loads(correct)
                except:
                    try:
                        correct = [int(x.strip()) for x in correct.split(",")]
                    except:
                        correct = []
            correct_set = set(correct)
            selected_indices = []
            for ans in user_ans:
                if ans in answers:
                    selected_indices.append(answers.index(ans) + 1)
            return set(selected_indices) == correct_set
        
        return False
    
    def get_all_questions():
        return []
    
    def get_all_users(role=None):
        return []
    
    def get_user_submissions(email):
        return []
    
    def get_all_submissions():
        return []

# Nhập các thư viện cho xuất file
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    # Hiển thị thông báo chỉ khi đang chạy trong Streamlit
    if 'streamlit' in sys.modules:
        st.warning("Module python-docx không được cài đặt. Tính năng xuất DOCX sẽ không hoạt động.")

# Sử dụng WD_ALIGN_PARAGRAPH nếu có thể, nếu không tạo class thay thế
try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    class WD_ALIGN_PARAGRAPH:
        CENTER = 1
        RIGHT = 2
        LEFT = 0

# Hỗ trợ xuất PDF với reportlab nếu cần
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    # Hiển thị thông báo chỉ khi đang chạy trong Streamlit
    if 'streamlit' in sys.modules:
        st.warning("Module reportlab không được cài đặt. Tính năng xuất PDF sẽ bị hạn chế.")

# Hàm kiểm tra cài đặt và phiên bản của FPDF
def check_fpdf_installed():
    try:
        # Kiểm tra phiên bản của fpdf
        fpdf_pkg = pkg_resources.get_distribution("fpdf")
        st.success(f"FPDF đã được cài đặt, phiên bản: {fpdf_pkg.version}")
        return True
    except pkg_resources.DistributionNotFound:
        st.error("FPDF chưa được cài đặt. Hãy cài đặt bằng lệnh: pip install fpdf2")
        return False
    except Exception as e:
        st.error(f"Lỗi khi kiểm tra FPDF: {str(e)}")
        return False

# Chuẩn bị font tiếng Việt
def setup_vietnamese_fonts():
    """Cài đặt và đăng ký font cho tiếng Việt"""
    current_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()
    
    # Các đường dẫn có thể chứa font
    font_dirs = [
        os.path.join(current_dir, 'assets', 'fonts'),
        os.path.join(current_dir, 'fonts'),
        os.path.join(current_dir, 'assets'),
        current_dir,
        '/usr/share/fonts/truetype',
        '/usr/share/fonts/truetype/dejavu',
        '/usr/share/fonts/TTF',
        'C:\\Windows\\Fonts',
    ]
    
    # Các font cần tìm
    font_files = [
        ('DejaVuSans', 'DejaVuSans.ttf'),
        ('DejaVuSans-Bold', 'DejaVuSans-Bold.ttf'),
        ('DejaVuSans-Oblique', 'DejaVuSans-Oblique.ttf'),
        ('Arial', 'arial.ttf'),
        ('Arial-Bold', 'arialbd.ttf'),
        ('Arial-Italic', 'ariali.ttf'),
    ]
    
    registered_fonts = []
    
    # Tìm và đăng ký font
    for font_name, font_file in font_files:
        for font_dir in font_dirs:
            font_path = os.path.join(font_dir, font_file)
            if os.path.exists(font_path):
                try:
                    # Đăng ký font với reportlab nếu cần
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                    except Exception as e:
                        print(f"Không thể đăng ký font {font_name} cho reportlab: {str(e)}")
                    
                    registered_fonts.append((font_name, font_path))
                    break
                except Exception as e:
                    print(f"Lỗi khi đăng ký font {font_name}: {str(e)}")
    
    return registered_fonts

# Đảm bảo có font DejaVu Unicode tại assets/fonts; nếu thiếu sẽ tải về
def ensure_dejavu_fonts():
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()
        fonts_dir = os.path.join(base_dir, 'assets', 'fonts')
        os.makedirs(fonts_dir, exist_ok=True)

        files = {
            'DejaVuSans.ttf': [
                'https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf',
                'https://raw.githubusercontent.com/dejavu-fonts/dejavu-fonts/master/ttf/DejaVuSans.ttf',
            ],
            'DejaVuSans-Bold.ttf': [
                'https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans-Bold.ttf',
                'https://raw.githubusercontent.com/dejavu-fonts/dejavu-fonts/master/ttf/DejaVuSans-Bold.ttf',
            ],
            'DejaVuSans-Oblique.ttf': [
                'https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans-Oblique.ttf',
                'https://raw.githubusercontent.com/dejavu-fonts/dejavu-fonts/master/ttf/DejaVuSans-Oblique.ttf',
            ],
        }

        local_paths = {}
        for fname, urls in files.items():
            local_path = os.path.join(fonts_dir, fname)
            local_paths[fname] = local_path
            if not os.path.exists(local_path):
                downloaded = False
                for url in urls:
                    try:
                        urllib.request.urlretrieve(url, local_path)
                        downloaded = True
                        break
                    except Exception as e:
                        continue
                if not downloaded:
                    print(f"Không thể tải font {fname} từ bất kỳ nguồn nào")
                    # Nếu một file tải thất bại, tiếp tục; sẽ fallback sau
        # Trả về dict các path đã có (có thể thiếu một vài file)
        return local_paths
    except Exception as e:
        print(f"Lỗi ensure_dejavu_fonts: {e}")
        return {}

def format_date(date_value):
    """Định dạng ngày tháng từ nhiều kiểu dữ liệu khác nhau"""
    if not date_value:
        return "N/A"
    
    try:
        # Nếu là số nguyên (timestamp)
        if isinstance(date_value, (int, float)):
            return datetime.fromtimestamp(date_value).strftime("%d/%m/%Y")
        
        # Nếu là chuỗi ISO (từ Supabase)
        elif isinstance(date_value, str):
            try:
                # Thử parse chuỗi ISO
                dt = datetime.fromisoformat(date_value.replace('Z', '+00:00'))
                return dt.strftime("%d/%m/%Y")
            except:
                # Nếu không phải ISO, trả về nguyên bản
                return date_value
        
        # Nếu đã là đối tượng datetime
        elif isinstance(date_value, datetime):
            return date_value.strftime("%d/%m/%Y")
            
        # Các trường hợp khác, trả về dạng chuỗi
        else:
            return str(date_value)
    except Exception as e:
        print(f"Error formatting date: {e}, value type: {type(date_value)}, value: {date_value}")
        return "N/A"

def get_buffer_content(buffer):
    """Lấy nội dung từ buffer một cách an toàn"""
    if buffer is None:
        return None
    
    try:
        # Đảm bảo buffer ở đầu
        current_pos = buffer.tell()
        buffer.seek(0)
        
        # Thử getvalue() trước
        content = buffer.getvalue()
        
        # Nếu getvalue() trả về None hoặc rỗng, thử read()
        if not content:
            content = buffer.read()
            buffer.seek(0)
        else:
            # Reset lại vị trí ban đầu nếu đã move
            buffer.seek(current_pos)
        
        return content
    except Exception as e:
        print(f"Lỗi khi đọc buffer: {e}")
        return None

def create_download_button(buffer, file_type, filename, button_text):
    """Tạo nút download sử dụng st.download_button"""
    try:
        # Lấy nội dung từ buffer
        content = get_buffer_content(buffer)
        
        if content is None or len(content) < 100:
            st.error(f"Không thể tạo file {file_type}: Buffer rỗng hoặc không hợp lệ")
            return False
        
        # Kiểm tra signature
        if file_type == "docx" and content[:2] != b'PK':
            st.error("File DOCX không hợp lệ (thiếu signature ZIP)")
            return False
        elif file_type == "pdf" and not content.startswith(b'%PDF'):
            st.error("File PDF không hợp lệ (thiếu signature %PDF)")
            return False
        
        # Xác định MIME type
        mime_types = {
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pdf": "application/pdf"
        }
        mime_type = mime_types.get(file_type, "application/octet-stream")
        
        # Tạo key duy nhất dựa trên filename, thời gian và hash của content
        import hashlib
        content_hash = hashlib.md5(content[:1000]).hexdigest()[:8]  # Lấy 8 ký tự đầu của hash
        unique_key = f"dl_{file_type}_{hashlib.md5(filename.encode()).hexdigest()[:8]}_{content_hash}"
        
        # Sử dụng st.download_button
        st.download_button(
            label=button_text,
            data=content,
            file_name=filename,
            mime=mime_type,
            key=unique_key
        )
        return True
    except Exception as e:
        print(f"Lỗi khi tạo download button {file_type}: {e}")
        import traceback
        traceback.print_exc()
        st.error(f"Lỗi khi tạo nút tải xuống: {str(e)}")
        return False

def get_download_link_docx(buffer, filename, text):
    """Tạo download button cho file DOCX - tương thích ngược"""
    # Sử dụng hàm mới với st.download_button
    return create_download_button(buffer, "docx", filename, text)

def get_download_link_pdf(buffer, filename, text):
    """Tạo download button cho file PDF - tương thích ngược"""
    # Sử dụng hàm mới với st.download_button
    return create_download_button(buffer, "pdf", filename, text)

def sanitize_sheet_name(name):
    """Làm sạch tên sheet để phù hợp với Excel (loại bỏ ký tự không hợp lệ và giới hạn độ dài)"""
    if not name:
        return "Sheet"
    
    # Loại bỏ các ký tự không hợp lệ trong Excel sheet name
    invalid_chars = [':', '\\', '/', '?', '*', '[', ']']
    cleaned = name
    for char in invalid_chars:
        cleaned = cleaned.replace(char, '-')
    
    # Giới hạn độ dài tối đa 31 ký tự (giới hạn của Excel)
    if len(cleaned) > 31:
        cleaned = cleaned[:31]
    
    # Loại bỏ khoảng trắng đầu/cuối
    cleaned = cleaned.strip()
    
    # Nếu sau khi làm sạch bị rỗng, dùng tên mặc định
    if not cleaned:
        cleaned = "Sheet"
    
    return cleaned

def export_to_excel(dataframes, sheet_names, filename, include_summary=True, questions=None, submissions=None):
    """Tạo file Excel với nhiều sheet từ các DataFrame, bao gồm phần tổng hợp điểm và tự động căn chỉnh"""
    try:
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            used_names = set()
            for df, sheet_name in zip(dataframes, sheet_names):
                # Làm sạch sheet name để tránh lỗi ký tự không hợp lệ
                base_name = sanitize_sheet_name(sheet_name)
                clean_sheet_name = base_name
                # Đảm bảo duy nhất trong workbook
                suffix = 1
                while clean_sheet_name in used_names:
                    candidate = f"{base_name[:28]}-{suffix}" if len(base_name) > 28 else f"{base_name}-{suffix}"
                    clean_sheet_name = sanitize_sheet_name(candidate)
                    suffix += 1
                used_names.add(clean_sheet_name)

                # Ghi DataFrame vào Excel
                df.to_excel(writer, sheet_name=clean_sheet_name, index=False)

                # Lấy worksheet để format
                worksheet = writer.sheets[clean_sheet_name]

                # Format header
                header_fill = PatternFill(start_color="E9E9E9", end_color="E9E9E9", fill_type="solid")
                header_font = Font(bold=True, size=11)
                header_alignment = Alignment(horizontal="center", vertical="center")

                for cell in worksheet[1]:
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = header_alignment

                # Tự động điều chỉnh độ rộng cột theo nội dung (giới hạn theo A4)
                # A4 width: ~210mm, với margin ~10mm mỗi bên = ~190mm usable
                # Approx: 1 character width = 0.1 inch = 2.54mm
                # Max columns on A4 portrait: ~75 characters total width
                max_width_chars = 75  # Giới hạn cho A4 portrait
                total_char_width = 0

                for idx, column in enumerate(worksheet.columns, 1):
                    column_letter = get_column_letter(idx)
                    max_length = 0
                    column_cells = list(column)

                    # Tìm độ dài tối đa trong cột
                    for cell in column_cells:
                        try:
                            if cell.value:
                                cell_str = str(cell.value)
                                # Giới hạn chiều dài cho tính toán
                                if len(cell_str) > 100:
                                    cell_str = cell_str[:100]
                                max_length = max(max_length, len(cell_str))
                        except:
                            pass

                    # Điều chỉnh độ rộng (min 10, max 50 để fit A4)
                    adjusted_width = min(max(max_length + 2, 10), 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
                    total_char_width += adjusted_width

                # Nếu tổng độ rộng vượt quá giới hạn, scale lại
                if total_char_width > max_width_chars:
                    scale_factor = max_width_chars / total_char_width
                    for idx in range(1, len(df.columns) + 1):
                        column_letter = get_column_letter(idx)
                        current_width = worksheet.column_dimensions[column_letter].width
                        worksheet.column_dimensions[column_letter].width = max(current_width * scale_factor, 8)

                # Điều chỉnh chiều cao hàng header
                worksheet.row_dimensions[1].height = 25

                # Thêm phần tổng hợp điểm nếu cần và có dữ liệu questions/submissions
                if include_summary and questions and submissions and "Tất cả bài nộp" in sheet_name:
                    # Tính tổng hợp điểm trắc nghiệm và tự luận
                    max_multiple_choice = sum([q.get("score", 0) for q in questions if q.get("type") in ["Checkbox", "Combobox"]])
                    max_essay = sum([q.get("score", 0) for q in questions if q.get("type") == "Essay"])
                    max_total = sum([q.get("score", 0) for q in questions])

                    total_multiple_choice = 0
                    total_essay = 0
                    total_score = 0
                    correct_answers = {}

                    # Tính từ các submissions
                    try:
                        for s in submissions:
                            responses = s.get("responses", {})
                            if isinstance(responses, str):
                                try:
                                    responses = json.loads(responses)
                                except:
                                    responses = {}
                            if not isinstance(responses, dict):
                                responses = {}

                            total_score += s.get("score", 0)

                            for q in questions:
                                q_id = str(q.get("id", ""))
                                q_type = q.get("type", "")
                                user_ans = responses.get(q_id, [])
                                if not isinstance(user_ans, list):
                                    user_ans = [user_ans] if user_ans is not None else []

                                try:
                                    from database_helper import check_answer_correctness as db_check_answer
                                    is_correct = db_check_answer(user_ans, q)
                                except ImportError:
                                    is_correct = check_answer_correctness(user_ans, q)

                                if is_correct:
                                    if q_id not in correct_answers:
                                        correct_answers[q_id] = 0
                                    correct_answers[q_id] += 1

                                    points = q.get("score", 0)
                                    if q_type in ["Checkbox", "Combobox"]:
                                        total_multiple_choice += points
                                    elif q_type == "Essay":
                                        total_essay += points
                    except Exception as calc_error:
                        print(f"Lỗi khi tính tổng hợp điểm trong Excel: {calc_error}")

                    # Thêm dòng trống
                    next_row = len(df) + 3
                    worksheet.cell(row=next_row, column=1).value = "TỔNG KẾT"
                    worksheet.cell(row=next_row, column=1).font = Font(bold=True, size=12)

                    # Thêm bảng tổng hợp
                    summary_data = [
                        ["Chỉ tiêu", "Giá trị"],
                        ["Điểm trắc nghiệm", f"{total_multiple_choice}/{max_multiple_choice}" if max_multiple_choice > 0 else "0/0"],
                        ["Điểm tự luận", f"{total_essay}/{max_essay}" if max_essay > 0 else "0/0"],
                        ["Tổng điểm", f"{total_score}/{max_total * len(submissions)}" if submissions else f"{total_score}/{max_total}"],
                        ["Tỷ lệ đúng trắc nghiệm", f"{(total_multiple_choice / (max_multiple_choice * len(submissions)) * 100):.1f}%" if max_multiple_choice > 0 and submissions else "0%"],
                        ["Tỷ lệ đúng tự luận", f"{(total_essay / (max_essay * len(submissions)) * 100):.1f}%" if max_essay > 0 and submissions else "0%"],
                    ]

                    start_row = next_row + 1
                    for i, row_data in enumerate(summary_data):
                        for j, value in enumerate(row_data, 1):
                            cell = worksheet.cell(row=start_row + i, column=j, value=value)
                            if i == 0:  # Header row
                                cell.fill = header_fill
                                cell.font = header_font
                                cell.alignment = header_alignment
                            else:
                                cell.alignment = Alignment(horizontal="left", vertical="center")

                                if j == 1:  # First column (labels)
                                    cell.font = Font(bold=True)

                                # Border
                                thin_border = Border(
                                    left=Side(style='thin'),
                                    right=Side(style='thin'),
                                    top=Side(style='thin'),
                                    bottom=Side(style='thin')
                                )
                                cell.border = thin_border

                    # Điều chỉnh độ rộng cột cho bảng tổng hợp
                    worksheet.column_dimensions[get_column_letter(1)].width = max(20, worksheet.column_dimensions[get_column_letter(1)].width)
                    worksheet.column_dimensions[get_column_letter(2)].width = max(20, worksheet.column_dimensions[get_column_letter(2)].width)

                # Đặt orientation và page size cho A4
                worksheet.page_setup.orientation = 'portrait'
                worksheet.page_setup.paperSize = worksheet.PAPERSIZE_A4

        # Đảm bảo dữ liệu được ghi
        output.flush()
        output.seek(0)

        # Lấy data
        data = output.getvalue()
        if not data or len(data) < 100:
            # Thử read nếu getvalue không có
            output.seek(0)
            data = output.read()
            output.seek(0)

        # Kiểm tra signature Excel (XLSX là ZIP format)
        if not data or len(data) < 100:
            raise ValueError(f"Excel buffer is empty or too small (length: {len(data) if data else 0})")

        if data[:2] != b'PK':
            raise ValueError("Excel buffer does not contain valid XLSX file (missing ZIP signature)")

        # Sử dụng st.download_button thay vì data URI
        output.seek(0)
        st.download_button(
            label=f"📥 Tải xuống {filename}",
            data=data,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"excel_{filename}_{id(output)}"
        )
        return True
    except Exception as e:
        print(f"Lỗi khi tạo Excel: {e}")
        import traceback
        traceback.print_exc()
        st.error(f"Không thể tạo file Excel: {str(e)}")
        return False

def dataframe_to_docx(df, title, filename):
    """Tạo file DOCX từ DataFrame"""
    try:
        doc = Document()
        
        # Thiết lập font chữ mặc định
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Thêm tiêu đề
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm thời gian xuất báo cáo
        time_paragraph = doc.add_paragraph(f"Thời gian xuất báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Xử lý trường hợp DataFrame rỗng hoặc không có cột để tránh DOCX hỏng
        if df is None or not hasattr(df, 'columns') or len(df.columns) == 0:
            doc.add_paragraph("Không có dữ liệu để hiển thị.")
        else:
            # Tạo bảng
            # Thêm một hàng cho tiêu đề cột
            table = doc.add_table(rows=1, cols=len(df.columns), style='Table Grid')
            
            # Thêm tiêu đề cột
            header_cells = table.rows[0].cells
            for i, col_name in enumerate(df.columns):
                header_cells[i].text = str(col_name)
                # Đặt kiểu cho tiêu đề
                for paragraph in header_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(str(col_name))
                    run.bold = True
            
            # Thêm dữ liệu
            if len(df) == 0:
                # Nếu không có dòng dữ liệu, thêm một dòng thông báo
                row_cells = table.add_row().cells
                if len(row_cells) > 0:
                    row_cells[0].text = "Không có dữ liệu"
                    for paragraph in row_cells[0].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
                        # Căn giữa cho các ô
                        for paragraph in row_cells[i].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm chân trang
        doc.add_paragraph()
        footer = doc.add_paragraph("Hệ thống Khảo sát & Đánh giá")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Lưu tệp vào buffer - đảm bảo cách xử lý đúng
        buffer = io.BytesIO()
        try:
            # Đảm bảo document được lưu hoàn toàn
            # Tạo buffer mới để lưu
            temp_buffer = io.BytesIO()
            doc.save(temp_buffer)
            
            # Đảm bảo dữ liệu được ghi hoàn toàn
            temp_buffer.flush()
            temp_buffer.seek(0)
            
            # Copy nội dung sang buffer chính
            buffer.write(temp_buffer.getvalue())
            temp_buffer.close()
            
            # Đảm bảo buffer ở đầu
            buffer.seek(0)
            
            # Kiểm tra nội dung
            content = buffer.getvalue()
            if not content or len(content) < 100:
                raise ValueError(f"DOCX buffer is empty or too small (length: {len(content) if content else 0})")
            
            # Kiểm tra signature DOCX (PK = ZIP format)
            if content[:2] != b'PK':
                raise ValueError("DOCX buffer does not contain valid DOCX file (missing ZIP signature)")
            
            # Đảm bảo buffer ở đầu để sẵn sàng đọc
            buffer.seek(0)
            return buffer
        except Exception as save_error:
            print(f"Lỗi khi lưu DOCX vào buffer: {save_error}")
            import traceback
            traceback.print_exc()
            # Đóng buffer và trả về None
            try:
                buffer.close()
            except:
                pass
            return None
            
    except Exception as e:
        print(f"Lỗi khi tạo DOCX: {str(e)}")
        import traceback
        traceback.print_exc()
        st.error(f"Không thể tạo file DOCX: {str(e)}")
        # Trả về None nếu lỗi
        return None

class UNIOCDF_FPDF(FPDF):
    """Lớp PDF tùy chỉnh hỗ trợ Unicode đầy đủ"""
    def __init__(self, orientation='P', unit='mm', format='A4', title='Báo cáo'):
        super().__init__(orientation=orientation, unit=unit, format=format)
        self.title = title
        
        # Khắc phục lỗi tiếng Việt bằng cách thiết lập encode utf8
        self.set_doc_option('core_fonts_encoding', 'utf-8')
        
        # Sử dụng DejaVu Sans (hỗ trợ Unicode) thay vì Times New Roman
        self.add_font('DejaVu', '', font_path='DejaVuSans.ttf', uni=True)
        self.add_font('DejaVu', 'B', font_path='DejaVuSans-Bold.ttf', uni=True)
        self.add_font('DejaVu', 'I', font_path='DejaVuSans-Oblique.ttf', uni=True)
        
    def header(self):
        # Font và tiêu đề
        self.set_font('DejaVu', 'B', 15)
        
        # Tiêu đề ở giữa
        self.cell(0, 10, self.title, 0, 1, 'C')
        
        # Thời gian
        self.set_font('DejaVu', 'I', 8)
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        self.cell(0, 5, f'Thời gian xuất báo cáo: {timestamp}', 0, 1, 'R')
        
        # Line break
        self.ln(5)
    
    def footer(self):
        # Vị trí cách đáy 15 mm
        self.set_y(-15)
        
        # Font
        self.set_font('DejaVu', 'I', 8)
        
        # Số trang
        self.cell(0, 10, f'Trang {self.page_no()}/{self.alias_nb_pages()}', 0, 0, 'C')
        
        # Thêm chân trang hệ thống
        self.cell(0, 10, 'Hệ thống Khảo sát & Đánh giá', 0, 0, 'R')

# Helper: set font an toàn với fallback
def _set_font_safe(pdf, style='', size=10):
    """Set font với fallback nếu font không tồn tại"""
    font_name = getattr(pdf, '_active_font_name', 'DejaVu')
    
    # Map style
    style_map = {
        'B': 'B',
        'I': 'I',
        'BI': 'BI',
        '': ''
    }
    font_style = style_map.get(style, '')
    
    try:
        pdf.set_font(font_name, font_style, size)
        return font_name
    except:
        try:
            pdf.set_font('Arial', font_style, size)
            return 'Arial'
        except:
            try:
                pdf.set_font('Helvetica', font_style, size)
                return 'Helvetica'
            except:
                pdf.set_font('Arial', '', size)  # Last resort - regular only
                return 'Arial'

# Helper: tính chiều cao cần thiết cho multi_cell với fpdf2
def _measure_multicell_height(pdf, width, text, line_height=5):
    try:
        lines = pdf.multi_cell(width, line_height, text, border=0, align='L', split_only=True)
        return max(line_height, len(lines) * line_height)
    except Exception:
        # Fallback ước lượng thô khi split_only không khả dụng
        avg_chars_per_line = max(1, int(width / max(1, pdf.get_string_width('W'))))
        num_lines = max(1, (len(str(text)) // avg_chars_per_line) + 1)
        return num_lines * line_height

# Tạo một instance FPDF có khả năng xử lý Unicode
def create_unicode_pdf(orientation='P', format='A4', title='Báo cáo'):
    """Tạo FPDF với hỗ trợ Unicode"""
    pdf = None
    font_name = None
    
    try:
        # Kiểm tra xem font đã được tìm thấy chưa
        font_dirs = [
            os.path.join(os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd(), 'assets', 'fonts'),
            os.path.join(os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd(), 'fonts'),
            os.path.join(os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd(), 'assets'),
            os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd(),
            '/usr/share/fonts/truetype/dejavu',
            '/usr/share/fonts/truetype',
            '/usr/share/fonts/TTF',
            'C:\\Windows\\Fonts',
        ]
        
        # Tìm font DejaVu Sans (bao gồm cả các biến thể Condensed nếu có)
        font_found = False
        font_paths = {
            'DejaVuSans.ttf': None,
            'DejaVuSans-Bold.ttf': None,
            'DejaVuSans-Oblique.ttf': None
        }
        alt_variants = [
            ('DejaVuSans.ttf', ['DejaVuSansCondensed.ttf', 'DejaVuSansCondensed.ttf']),
            ('DejaVuSans-Bold.ttf', ['DejaVuSansCondensed-Bold.ttf', 'DejaVuSans-Bold.ttf']),
            ('DejaVuSans-Oblique.ttf', ['DejaVuSansCondensed-Oblique.ttf', 'DejaVuSans-Oblique.ttf'])
        ]
        
        # Tìm từng font file (ưu tiên bản tiêu chuẩn, nếu không có thử các biến thể thay thế)
        for font_dir in font_dirs:
            for font_file in list(font_paths.keys()):
                if font_paths[font_file]:
                    continue
                # Thử exact match
                font_path = os.path.join(font_dir, font_file)
                if os.path.exists(font_path):
                    font_paths[font_file] = font_path
                    continue
                # Thử các biến thể thay thế
                for alt in [a for key, alts in alt_variants if key == font_file for a in alts]:
                    alt_path = os.path.join(font_dir, alt)
                    if os.path.exists(alt_path):
                        font_paths[font_file] = alt_path
                        break
        
        # Kiểm tra xem có đủ cả 3 font không; nếu thiếu, thử tải về assets/fonts
        if not all(font_paths.values()):
            downloaded = ensure_dejavu_fonts()
            for k in list(font_paths.keys()):
                if not font_paths[k]:
                    candidate = downloaded.get(k)
                    if candidate and os.path.exists(candidate):
                        font_paths[k] = candidate
        if all(font_paths.values()):
            font_found = True
        
        # Tạo PDF mới
        pdf = FPDF(orientation=orientation, unit='mm', format=format)
        
        # Thiết lập mã hóa UTF-8
        pdf.set_doc_option('core_fonts_encoding', 'utf-8')
        
        # Thêm các font Unicode nếu tìm thấy (cho phép thiếu style: dùng regular cho B/I nếu cần)
        if font_found:
            try:
                regular_path = font_paths['DejaVuSans.ttf']
                bold_path = font_paths.get('DejaVuSans-Bold.ttf') or regular_path
                italic_path = font_paths.get('DejaVuSans-Oblique.ttf') or regular_path
                pdf.add_font('DejaVu', '', regular_path, uni=True)
                pdf.add_font('DejaVu', 'B', bold_path, uni=True)
                pdf.add_font('DejaVu', 'I', italic_path, uni=True)
                font_name = 'DejaVu'
            except Exception as font_error:
                print(f"Lỗi khi thêm font DejaVu: {font_error}")
                font_found = False
        
        # Nếu không tìm thấy DejaVu, thử tìm font khác hoặc dùng built-in
        if not font_found:
            # Thử tìm Arial hoặc font hệ thống khác
            arial_paths = ['arial.ttf', 'Arial.ttf', 'Arial.TTF']
            arial_found = False
            
            for font_dir in font_dirs:
                for arial_name in arial_paths:
                    arial_path = os.path.join(font_dir, arial_name)
                    if os.path.exists(arial_path):
                        try:
                            pdf.add_font('Arial', '', arial_path, uni=True)
                            pdf.add_font('Arial', 'B', arial_path, uni=True)  # Dùng cùng font cho bold
                            pdf.add_font('Arial', 'I', arial_path, uni=True)  # Dùng cùng font cho italic
                            font_name = 'Arial'
                            arial_found = True
                            break
                        except Exception as arial_error:
                            print(f"Lỗi khi thêm font Arial: {arial_error}")
                            continue
                    if arial_found:
                        break
                if arial_found:
                    break
            
            # Nếu vẫn không tìm thấy font Unicode, sử dụng built-in fonts (không hỗ trợ Unicode tốt)
            if not arial_found:
                font_name = 'Arial'  # FPDF built-in font (hạn chế Unicode)
                print("Cảnh báo: Sử dụng built-in fonts, có thể không hiển thị đúng ký tự tiếng Việt")
        
        # Thiết lập các tùy chọn khác
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.alias_nb_pages()
        
        # Lưu tên font đã sử dụng vào PDF object để dùng sau
        pdf._active_font_name = font_name
        
        return pdf
        
    except Exception as e:
        print(f"Lỗi tạo PDF: {str(e)}")
        traceback.print_exc()
        
        # Phương án dự phòng - sử dụng FPDF cơ bản với built-in fonts
        try:
            pdf = FPDF(orientation=orientation, format=format)
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.alias_nb_pages()
            pdf._active_font_name = 'Arial'  # Built-in font
            return pdf
        except Exception as e2:
            print(f"Lỗi khi tạo PDF dự phòng: {str(e2)}")
            return None

def dataframe_to_pdf_reportlab(df, title, filename):
    """Tạo file PDF từ DataFrame sử dụng ReportLab (tránh font issues)"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import mm
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                                rightMargin=10*mm, leftMargin=10*mm,
                                topMargin=15*mm, bottomMargin=15*mm)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#000000'),
            alignment=1,  # Center
            spaceAfter=12
        )
        
        data_style = ParagraphStyle(
            'DataStyle',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#000000'),
            leading=10
        )
        
        # Title
        story = [Paragraph(title, title_style)]
        story.append(Spacer(1, 6*mm))
        
        # Timestamp
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        timestamp_para = Paragraph(f"Thời gian xuất báo cáo: {timestamp}", styles['Normal'])
        story.append(timestamp_para)
        story.append(Spacer(1, 4*mm))
        
        # Prepare table data (xử lý df rỗng an toàn)
        table_data = []
        if df is None or not hasattr(df, 'columns') or len(df.columns) == 0:
            table_data = [["Không có dữ liệu"]]
        else:
            # Header - sử dụng Paragraph để tránh overlap
            header = [Paragraph(str(col), data_style) for col in df.columns]
            table_data.append(header)
            
            # Data rows (limit để tránh file quá lớn)
            max_rows = min(500, len(df))
            if max_rows == 0:
                empty_row = [Paragraph("Không có dữ liệu", data_style)] + [Paragraph("", data_style)] * (len(df.columns) - 1)
                table_data.append(empty_row)
            else:
                for i in range(max_rows):
                    row = []
                    for j in range(len(df.columns)):
                        cell_value = str(df.iloc[i, j])
                        # Giới hạn độ dài để tránh overflow
                        if len(cell_value) > 100:
                            cell_value = cell_value[:97] + "..."
                        row.append(Paragraph(cell_value, data_style))
                    table_data.append(row)
        
        # Create table với độ rộng cột cố định để tránh overlap
        num_cols = len(df.columns) if df is not None and hasattr(df, 'columns') and len(df.columns) > 0 else 1
        
        # Tính toán độ rộng cột cố định cho A4 (trừ margin)
        available_width = A4[0] - 20*mm  # Trừ margin trái phải
        col_widths = []
        
        if num_cols == 1:
            col_widths = [available_width]
        elif num_cols == 2:
            col_widths = [available_width * 0.6, available_width * 0.4]
        elif num_cols == 3:
            col_widths = [available_width * 0.4, available_width * 0.3, available_width * 0.3]
        elif num_cols == 4:
            col_widths = [available_width * 0.3, available_width * 0.25, available_width * 0.25, available_width * 0.2]
        elif num_cols == 5:
            col_widths = [available_width * 0.25, available_width * 0.2, available_width * 0.2, available_width * 0.2, available_width * 0.15]
        else:
            # Nhiều hơn 5 cột - chia đều
            col_widths = [available_width / num_cols] * num_cols
        
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        
        # Style table - tối ưu cho Paragraph objects
        table.setStyle(TableStyle([
            # Header
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E9E9E9')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            ('LEFTPADDING', (0, 0), (-1, 0), 4),
            ('RIGHTPADDING', (0, 0), (-1, 0), 4),
            
            # Body
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FFFFFF')),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#000000')),
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),  # Left align cho Paragraph
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Top align để tránh overlap
            ('LEFTPADDING', (0, 1), (-1, -1), 4),
            ('RIGHTPADDING', (0, 1), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
            ('TOPPADDING', (0, 1), (-1, -1), 4),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F9F9F9')]),
        ]))
        
        story.append(table)
        
        # Build PDF
        doc.build(story)
        
        # Validate buffer
        buffer.seek(0)
        content = buffer.getvalue()
        if not content or len(content) < 100:
            raise ValueError(f"PDF buffer is empty or too small (length: {len(content) if content else 0})")
        
        if not content.startswith(b'%PDF'):
            raise ValueError("PDF buffer does not contain valid PDF file (missing %PDF signature)")
        
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"Lỗi khi tạo PDF bằng ReportLab: {e}")
        import traceback
        traceback.print_exc()
        # Fallback to FPDF nếu ReportLab fail
        return dataframe_to_pdf_fpdf(df, title, filename)

def dataframe_to_pdf_fpdf(df, title, filename):
    """Tạo file PDF từ DataFrame sử dụng FPDF2 với hỗ trợ Unicode"""
    # Buffer sẽ được tạo khi lưu PDF, không tạo ở đây
    try:
        # Xác định hướng trang dựa vào số lượng cột
        orientation = 'L' if len(df.columns) > 5 else 'P'
        
        # Sử dụng FPDF2 có hỗ trợ Unicode
        pdf = create_unicode_pdf(orientation=orientation, title=title)
        
        if pdf is None:
            raise Exception("Không thể tạo đối tượng PDF")
        
        pdf.add_page()
        
        # Lấy tên font đã được add (nếu có)
        font_name = getattr(pdf, '_active_font_name', 'DejaVu')
        
        # Thêm tiêu đề - sử dụng font đã được add
        try:
            pdf.set_font(font_name, 'B', 16)
        except:
            try:
                pdf.set_font('Arial', 'B', 16)  # Fallback to built-in
                font_name = 'Arial'
            except:
                pdf.set_font('Helvetica', 'B', 16)  # Last resort
                font_name = 'Helvetica'
        
        pdf.cell(0, 10, title, 0, 1, 'C')
        
        # Thêm thời gian báo cáo
        try:
            pdf.set_font(font_name, 'I', 10)
        except:
            pdf.set_font(font_name, '', 10)  # Nếu không có italic, dùng regular
            
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        pdf.cell(0, 5, f'Thời gian xuất báo cao: {timestamp}', 0, 1, 'R')
        pdf.ln(5)
        
        # Xác định kích thước trang và số cột
        page_width = 297 if orientation == 'L' else 210
        margin = 10
        usable_width = page_width - 2*margin
        
        # Tính toán độ rộng cột hợp lý
        col_widths = []
        max_content_widths = []
        
        # Mặc định font cho nội dung
        try:
            pdf.set_font(font_name, '', 9)
        except:
            pdf.set_font('Arial', '', 9)
        
        # Ước tính độ rộng tối đa cho mỗi cột
        for i, col in enumerate(df.columns):
            # Độ rộng tiêu đề
            header_width = pdf.get_string_width(str(col)) + 6  # Thêm padding
            
            # Độ rộng nội dung (kiểm tra 20 dòng đầu tiên)
            content_widths = []
            for j in range(min(20, len(df))):
                cell_content = str(df.iloc[j, i])
                # Giới hạn độ dài chuỗi để tránh quá rộng
                if len(cell_content) > 100:
                    cell_content = cell_content[:97] + "..."
                content_width = pdf.get_string_width(cell_content) + 6  # Thêm padding
                content_widths.append(content_width)
            
            max_content_width = max(content_widths) if content_widths else 0
            max_width = max(header_width, max_content_width)
            
            # Giới hạn độ rộng cột
            max_col_width = 70  # tăng trần để tránh lỗi không đủ chỗ
            # đảm bảo tối thiểu vừa 1 ký tự 'W' + padding
            min_char = pdf.get_string_width('W') + 4
            col_width = min(max_col_width, max(min_char, max_width))
            
            col_widths.append(col_width)
            max_content_widths.append(max_content_width)
        
        # Điều chỉnh để tổng độ rộng không vượt quá chiều rộng khả dụng
        total_width = sum(col_widths)
        
        # Đảm bảo độ rộng tối thiểu cho mỗi cột (đủ để hiển thị 1 ký tự + padding)
        min_width_per_col = pdf.get_string_width('W') + 6  # padding cho border
        min_total_width = min_width_per_col * len(df.columns)
        
        # Nếu tổng độ rộng tối thiểu vượt quá usable_width, cần điều chỉnh
        if min_total_width > usable_width:
            # Trường hợp quá nhiều cột: scale min_width xuống
            min_width_per_col = (usable_width / len(df.columns)) - 2  # trừ margin
            if min_width_per_col < 5:  # Không thể nhỏ hơn 5mm
                min_width_per_col = 5
            # Đặt lại tất cả cột về min_width
            col_widths = [min_width_per_col] * len(df.columns)
        elif total_width > usable_width:
            # Scale xuống nhưng đảm bảo không nhỏ hơn min_width_per_col
            scale_factor = usable_width / total_width
            col_widths = [max(min_width_per_col, width * scale_factor) for width in col_widths]
            # Điều chỉnh lại nếu vẫn vượt (do đảm bảo min_width)
            new_total = sum(col_widths)
            if new_total > usable_width:
                # Scale lại lần nữa
                scale_factor = usable_width / new_total
            col_widths = [width * scale_factor for width in col_widths]
        
        # Đảm bảo mỗi cột >= min_width cuối cùng (trước khi vẽ)
        for idx in range(len(col_widths)):
            if col_widths[idx] < min_width_per_col:
                col_widths[idx] = min_width_per_col
        
        # Tính lại tổng độ rộng sau khi đảm bảo min_width
        final_total = sum(col_widths)
        if final_total > usable_width:
            # Scale lại toàn bộ để vừa trang, nhưng chỉ scale một lần
            scale_factor = usable_width / final_total
            for idx in range(len(col_widths)):
                col_widths[idx] = col_widths[idx] * scale_factor
            
            # Kiểm tra lại sau scale - nếu vẫn có cột < min_width (do scale quá nhiều)
            # thì chỉ đặt lại các cột quá nhỏ thành min_width, chấp nhận có thể vượt trang một chút
            for idx in range(len(col_widths)):
                if col_widths[idx] < min_width_per_col:
                    col_widths[idx] = min_width_per_col
        
        # Tạo tiêu đề cột
        _set_font_safe(pdf, 'B', 10)
        pdf.set_fill_color(240, 240, 240)
        
        # Lưu vị trí bắt đầu
        start_x = pdf.get_x()
        start_y = pdf.get_y()
        
        # Vẽ header với độ cao đồng nhất
        header_height = 10
        for i, col_name in enumerate(df.columns):
            # Tính toán vị trí x mới
            new_x = start_x + sum(col_widths[:i])
            pdf.set_xy(new_x, start_y)
            
            # Cắt ngắn tên cột nếu quá dài và convert sang ASCII-safe nếu cần
            col_name_str = str(col_name)
            if len(col_name_str) > 25:
                col_name_str = col_name_str[:22] + "..."
            
            # Convert sang ASCII-safe để tránh font issue
            try:
                # Thử hiển thị nguyên gốc trước
                pdf.cell(col_widths[i], header_height, col_name_str, 1, 0, 'C', 1)
            except:
                # Fallback: convert sang ASCII-safe
                ascii_name = ''.join(c if ord(c) < 128 else '?' for c in col_name_str)
                pdf.cell(col_widths[i], header_height, ascii_name[:25], 1, 0, 'C', 1)
        
        pdf.ln(header_height)
        
        # Vẽ nội dung với font nhỏ hơn
        _set_font_safe(pdf, '', 8)
        
        # Chiều cao dòng cơ bản
        row_height = 7
        
        # Giới hạn số lượng hàng để tránh file quá lớn
        max_rows = min(1000, len(df))
        
        for i in range(max_rows):
            # Reset về đầu dòng mới
            row_start_x = pdf.get_x()
            row_start_y = pdf.get_y()
            max_height = row_height
            
            # Kiểm tra xem có đủ không gian cho dòng mới không
            if row_start_y + row_height > pdf.page_break_trigger:
                pdf.add_page()
                row_start_y = pdf.get_y()
            
            # Vẽ từng ô trong hàng hiện tại
            # Đảm bảo tất cả col_widths >= min_width_per_col trước khi tính cell_x
            min_w = pdf.get_string_width('W') + 6
            min_safe_width = max(min_width_per_col, min_w, 9)  # Tối thiểu 9mm
            
            # Fix tất cả col_widths trước, đảm bảo không update trong vòng lặp
            for idx in range(len(df.columns)):
                if col_widths[idx] < min_safe_width:
                    col_widths[idx] = min_safe_width
            
            for j, col_name in enumerate(df.columns):
                cell_x = row_start_x + sum(col_widths[:j])
                content = str(df.iloc[i, j])
                
                # col_widths[j] đã được đảm bảo >= min_safe_width ở trên
                # Cắt ngắn nội dung nếu quá dài để tránh vấn đề hiển thị
                if len(content) > 200:
                    content = content[:197] + "..."
                
                # Convert sang ASCII-safe nếu font không hỗ trợ Unicode tốt
                font_name = getattr(pdf, '_active_font_name', 'DejaVu')
                if font_name not in ['DejaVu', 'Arial']:  # Built-in fonts không hỗ trợ Unicode tốt
                    # Convert sang ASCII-safe
                    content = ''.join(c if ord(c) < 128 else '?' for c in content)
                
                # Tính số dòng cần thiết cho nội dung này
                try:
                    content_width = pdf.get_string_width(content)
                except:
                    # Fallback: ước tính thô
                    content_width = len(content) * pdf.get_string_width('W') / 10
                
                available_width = col_widths[j] - 6  # Trừ padding và border
                
                # Đảm bảo available_width luôn > 0 (col_widths[j] đã >= 9mm)
                if available_width <= 0:
                    available_width = 3  # Tối thiểu 3mm (do col_widths[j] >= 9mm nên available >= 3)
                
                if content_width > available_width:
                    # Ước tính số dòng cần thiết
                    num_lines = max(1, int(content_width / available_width) + 1)
                    # Tính chiều cao cần thiết
                    cell_height = max(row_height, num_lines * 5)  # 5mm cho mỗi dòng
                else:
                    cell_height = row_height
                
                # Cập nhật chiều cao tối đa cho dòng hiện tại
                max_height = max(max_height, cell_height)
                
                # Vẽ ô với nội dung - sử dụng col_widths[j] đã được đảm bảo
                pdf.set_xy(cell_x, row_start_y)
                # col_widths[j] đã được đảm bảo >= 9mm ở trên
                final_width = col_widths[j]
                try:
                    pdf.multi_cell(final_width, cell_height, content, 1, 'L')
                except Exception as e:
                    # Nếu vẫn lỗi, thử với độ rộng lớn hơn
                    final_width = max(final_width, 10)
                    try:
                        pdf.multi_cell(final_width, cell_height, content[:50] if len(content) > 50 else content, 1, 'L')
                    except:
                        # Last resort: chỉ render ký tự ASCII
                        ascii_content = ''.join(c for c in content if ord(c) < 128)[:50]
                        pdf.multi_cell(final_width, cell_height, ascii_content or "N/A", 1, 'L')
            
            # Di chuyển đến dòng tiếp theo
            pdf.set_y(row_start_y + max_height)
        
        # Lưu PDF vào buffer - luôn sử dụng dest='S' để lấy bytes
        buffer = io.BytesIO()  # Đảm bảo buffer sạch
        try:
            # Sử dụng dest='S' để lấy bytes trực tiếp (cách an toàn nhất)
            pdf_bytes = pdf.output(dest='S')
            
            # Xử lý các kiểu dữ liệu khác nhau
            if isinstance(pdf_bytes, bytes):
                buffer.write(pdf_bytes)
            elif isinstance(pdf_bytes, str):
                # Nếu là string, encode sang bytes
                buffer.write(pdf_bytes.encode('latin-1', errors='ignore'))
            elif hasattr(pdf_bytes, '__iter__'):
                # Nếu là iterable, chuyển thành bytes
                buffer.write(bytes(pdf_bytes))
            else:
                # Fallback: thử encode sang string rồi bytes
                buffer.write(str(pdf_bytes).encode('latin-1', errors='ignore'))
            
            buffer.flush()
            
        except Exception as output_error:
            print(f"Lỗi khi output PDF: {output_error}")
            # Thử lại với dest=buffer
            try:
                buffer = io.BytesIO()
                pdf.output(dest=buffer)
                buffer.flush()
            except Exception:
                raise ValueError(f"Không thể tạo PDF output: {str(output_error)}")
        
        # Đảm bảo buffer ở đầu và kiểm tra dữ liệu
        buffer.seek(0)
        content = buffer.getvalue()
        
        # Kiểm tra signature PDF
        if not content or len(content) < 100:
            # Thử đọc lại bằng read()
            buffer.seek(0)
            content = buffer.read()
            buffer.seek(0)
            
            if not content or len(content) < 100:
                raise ValueError(f"PDF buffer is empty or too small (length: {len(content) if content else 0}). Có thể PDF chưa được tạo đúng cách.")
        
        if not content.startswith(b'%PDF'):
            raise ValueError(f"PDF buffer does not contain valid PDF (missing %PDF signature). Buffer đầu: {content[:50] if content else 'None'}")
        
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"Lỗi khi tạo PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Tạo báo cáo đơn giản nếu gặp lỗi
        try:
            buffer = io.BytesIO()  # Tạo buffer mới
            simple_pdf = FPDF()
            simple_pdf.add_page()
            simple_pdf.set_font('Arial', 'B', 16)
            simple_pdf.cell(0, 10, title, 0, 1, 'C')
            simple_pdf.set_font('Arial', '', 10)
            error_msg = f'Khong the tao bao cao chi tiet. Vui long su dung dinh dang DOCX hoac Excel.\nLoi: {str(e)[:100]}'
            simple_pdf.multi_cell(0, 10, error_msg, 0, 'L')
            
            try:
                pdf_bytes = simple_pdf.output(dest='S')
                if isinstance(pdf_bytes, str):
                    pdf_bytes = pdf_bytes.encode('latin-1')
                buffer.write(pdf_bytes)
                buffer.flush()
            except Exception:
                simple_pdf.output(dest=buffer)
                buffer.flush()
            
            buffer.seek(0)
            
            # Kiểm tra lại
            content = buffer.getvalue()
            if content and len(content) > 100 and content.startswith(b'%PDF'):
                return buffer
            else:
                return None
        except Exception as e2:
            print(f"Không thể tạo báo cáo thay thế: {str(e2)}")
            return None

def create_student_report_pdf_reportlab(student_name, student_email, student_class, submission, questions, max_possible):
    """Tạo báo cáo chi tiết học viên dạng PDF sử dụng ReportLab (tránh font issues)"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.units import mm
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=10*mm, leftMargin=10*mm,
                                topMargin=15*mm, bottomMargin=15*mm)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#000000'),
            alignment=1,  # Center
            spaceAfter=12
        )
        
        # Story để chứa nội dung
        story = []
        
        # Title
        title = f"Báo cáo chi tiết - {student_name}"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 6*mm))
        
        # Timestamp
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        story.append(Paragraph(f"Thời gian xuất báo cáo: {timestamp}", styles['Normal']))
        story.append(Spacer(1, 4*mm))
        
        # Thông tin học viên
        story.append(Paragraph("<b>Thông tin học viên</b>", styles['Heading2']))
        info_data = [
            ["Thông tin", "Giá trị"],
            ["Họ và tên", student_name],
            ["Email", student_email],
            ["Lớp", student_class],
        ]
        
        # Xử lý timestamp
        submission_time = "Không xác định"
        if isinstance(submission.get("timestamp"), (int, float)):
            try:
                submission_time = datetime.fromtimestamp(submission.get("timestamp")).strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
        else:
            try:
                dt = datetime.fromisoformat(submission.get("timestamp", "").replace("Z", "+00:00"))
                submission_time = dt.strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
        
        info_data.append(["Thời gian nộp", submission_time])
        
        info_table = Table(info_data, colWidths=[50*mm, 140*mm])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E9E9E9')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FFFFFF')),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#000000')),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 1), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F9F9F9')]),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 8*mm))
        
        # Chi tiết câu trả lời
        story.append(Paragraph("<b>Chi tiết câu trả lời</b>", styles['Heading2']))
        
        # Đảm bảo responses đúng định dạng
        responses = submission.get("responses", {})
        if isinstance(responses, str):
            try:
                responses = json.loads(responses)
            except:
                responses = {}
        if not isinstance(responses, dict):
            responses = {}
        
        # Validate questions
        normalized_questions = []
        for q in questions:
            if not isinstance(q, dict):
                continue
            if isinstance(q.get("answers"), str):
                try:
                    q["answers"] = json.loads(q["answers"])
                except:
                    q["answers"] = [q["answers"]] if q.get("answers") else []
            if isinstance(q.get("correct"), str):
                try:
                    q["correct"] = json.loads(q["correct"])
                except:
                    try:
                        q["correct"] = [int(x.strip()) for x in q["correct"].split(",")]
                    except:
                        q["correct"] = []
            if not isinstance(q.get("answers"), list):
                q["answers"] = []
            if not isinstance(q.get("correct"), list):
                q["correct"] = []
            normalized_questions.append(q)
        questions = normalized_questions
        
        # Tính toán điểm
        total_correct = 0
        calculated_total_score = 0
        multiple_choice_score = 0
        essay_score = 0
        
        # Tạo bảng chi tiết với Paragraph objects để tránh overlap
        detail_data = []
        
        # Header với Paragraph objects
        header_style = ParagraphStyle(
            'HeaderStyle',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#000000'),
            alignment=1,  # Center
            fontName='Helvetica-Bold'
        )
        
        detail_data.append([
            Paragraph("Câu hỏi", header_style),
            Paragraph("Đáp án học viên", header_style),
            Paragraph("Đáp án đúng", header_style),
            Paragraph("Kết quả", header_style),
            Paragraph("Điểm", header_style)
        ])
        
        # Data style cho nội dung
        data_style = ParagraphStyle(
            'DataStyle',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#000000'),
            leading=10,
            fontName='Helvetica'
        )
        
        for q in questions:
            q_id = str(q.get("id", ""))
            user_ans = responses.get(q_id, [])
            if not isinstance(user_ans, list):
                user_ans = [user_ans] if user_ans is not None else []
            
            try:
                from database_helper import check_answer_correctness as db_check_answer
                is_correct = db_check_answer(user_ans, q)
            except ImportError:
                is_correct = check_answer_correctness(user_ans, q)
            
            q_type = q.get("type", "")
            if is_correct:
                total_correct += 1
                points = q.get("score", 0)
            else:
                points = 0
            
            if q_type in ["Checkbox", "Combobox"]:
                if is_correct:
                    multiple_choice_score += points
            elif q_type == "Essay":
                if is_correct:
                    essay_score += points
            
            calculated_total_score += points
            
            # Chuẩn bị dữ liệu cho bảng với Paragraph objects
            question_text = f"Câu {q.get('id', '')}: {q.get('question', '')}"
            
            if q_type == "Essay":
                user_answer_text = user_ans[0] if user_ans else "Không trả lời"
                correct_answer_text = "Câu hỏi tự luận"
                result = "Đã trả lời" if is_correct else "Không trả lời"
            else:
                user_answer_text = ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời"
                q_correct = q.get("correct", [])
                q_answers = q.get("answers", [])
                try:
                    expected = [q_answers[i - 1] for i in q_correct] if q_correct else []
                    correct_answer_text = ", ".join([str(a) for a in expected]) if expected else "Không có đáp án"
                except:
                    correct_answer_text = "Lỗi đáp án"
                result = "Đúng" if is_correct else "Sai"
            
            # Tạo Paragraph objects cho mỗi cell
            detail_data.append([
                Paragraph(question_text, data_style),
                Paragraph(user_answer_text, data_style),
                Paragraph(correct_answer_text, data_style),
                Paragraph(result, data_style),
                Paragraph(str(points), data_style)
            ])
        
        # Tạo bảng với độ rộng cột cố định để tránh overlap
        col_widths = [60*mm, 45*mm, 45*mm, 20*mm, 20*mm]
        detail_table = Table(detail_data, colWidths=col_widths, repeatRows=1)
        detail_table.setStyle(TableStyle([
            # Header
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E9E9E9')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            ('LEFTPADDING', (0, 0), (-1, 0), 4),
            ('RIGHTPADDING', (0, 0), (-1, 0), 4),
            
            # Body - tối ưu cho Paragraph objects
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FFFFFF')),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#000000')),
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),  # Left align cho Paragraph
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Top align để tránh overlap
            ('LEFTPADDING', (0, 1), (-1, -1), 4),
            ('RIGHTPADDING', (0, 1), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
            ('TOPPADDING', (0, 1), (-1, -1), 4),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F9F9F9')]),
        ]))
        story.append(detail_table)
        story.append(Spacer(1, 8*mm))
        
        # Tổng kết
        max_multiple_choice = sum([q.get("score", 0) for q in questions if q.get("type") in ["Checkbox", "Combobox"]])
        max_essay = sum([q.get("score", 0) for q in questions if q.get("type") == "Essay"])
        
        story.append(Paragraph("<b>Tổng kết</b>", styles['Heading2']))
        
        # Tạo bảng tổng kết với Paragraph objects
        summary_data = []
        
        # Header
        summary_header_style = ParagraphStyle(
            'SummaryHeaderStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#000000'),
            alignment=1,  # Center
            fontName='Helvetica-Bold'
        )
        
        summary_data.append([
            Paragraph("Chỉ tiêu", summary_header_style),
            Paragraph("Giá trị", summary_header_style)
        ])
        
        # Data rows với Paragraph objects
        summary_data_style = ParagraphStyle(
            'SummaryDataStyle',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#000000'),
            leading=12,
            fontName='Helvetica'
        )
        
        summary_data.append([
            Paragraph("Số câu đúng", summary_data_style),
            Paragraph(f"{total_correct}/{len(questions)}", summary_data_style)
        ])
        summary_data.append([
            Paragraph("Điểm trắc nghiệm", summary_data_style),
            Paragraph(f"{multiple_choice_score}/{max_multiple_choice}" if max_multiple_choice > 0 else "0/0", summary_data_style)
        ])
        summary_data.append([
            Paragraph("Điểm tự luận", summary_data_style),
            Paragraph(f"{essay_score}/{max_essay}" if max_essay > 0 else "0/0", summary_data_style)
        ])
        summary_data.append([
            Paragraph("Tổng điểm", summary_data_style),
            Paragraph(f"{calculated_total_score}/{max_possible}", summary_data_style)
        ])
        summary_data.append([
            Paragraph("Tỷ lệ đúng", summary_data_style),
            Paragraph(f"{(total_correct/len(questions)*100):.1f}%" if len(questions) > 0 else "0%", summary_data_style)
        ])
        summary_data.append([
            Paragraph("Tỷ lệ điểm", summary_data_style),
            Paragraph(f"{(calculated_total_score/max_possible*100):.1f}%" if max_possible > 0 else "0%", summary_data_style)
        ])
        
        summary_table = Table(summary_data, colWidths=[80*mm, 110*mm])
        summary_table.setStyle(TableStyle([
            # Header
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E9E9E9')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            ('LEFTPADDING', (0, 0), (-1, 0), 4),
            ('RIGHTPADDING', (0, 0), (-1, 0), 4),
            
            # Body - tối ưu cho Paragraph objects
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FFFFFF')),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#000000')),
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),  # Left align cho Paragraph
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Top align để tránh overlap
            ('LEFTPADDING', (0, 1), (-1, -1), 4),
            ('RIGHTPADDING', (0, 1), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
            ('TOPPADDING', (0, 1), (-1, -1), 4),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F9F9F9')]),
        ]))
        story.append(summary_table)
        
        # Build PDF
        doc.build(story)
        
        # Validate buffer
        buffer.seek(0)
        content = buffer.getvalue()
        if not content or len(content) < 100:
            raise ValueError(f"PDF buffer is empty or too small (length: {len(content) if content else 0})")
        if not content.startswith(b'%PDF'):
            raise ValueError("PDF buffer does not contain valid PDF file (missing %PDF signature)")
        
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"Lỗi khi tạo PDF bằng ReportLab: {e}")
        import traceback
        traceback.print_exc()
        # Fallback to FPDF nếu ReportLab fail
        return create_student_report_pdf_fpdf(student_name, student_email, student_class, submission, questions, max_possible)

def create_student_report_docx(student_name, student_email, student_class, submission, questions, max_possible):
    """Tạo báo cáo chi tiết bài làm của học viên dạng DOCX"""
    try:
        doc = Document()
        
        # Thiết lập font chữ mặc định là Times New Roman
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Thêm tiêu đề - font chữ hỗ trợ Unicode
        heading = doc.add_heading(f"Báo cáo chi tiết - {student_name}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm thời gian xuất báo cáo
        time_paragraph = doc.add_paragraph(f"Thời gian xuất báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Thêm thông tin học viên
        doc.add_heading("Thông tin học viên", level=2)
        info_table = doc.add_table(rows=4, cols=2, style='Table Grid')
        
        # Đặt độ rộng cột
        for cell in info_table.columns[0].cells:
            cell.width = Inches(1.5)
        for cell in info_table.columns[1].cells:
            cell.width = Inches(4.5)
        
        # Thiết lập màu nền cho hàng tiêu đề
        for i in range(4):
            # Sửa lỗi: Đảm bảo có runs trước khi truy cập
            cell = info_table.rows[i].cells[0]
            cell_paragraph = cell.paragraphs[0]
            if not cell_paragraph.runs:
                cell_paragraph.add_run(cell.text if cell.text else '')
            cell_paragraph.runs[0].font.bold = True
            
            # Thêm màu nền
            shading_elm = parse_xml(r'<w:shd {} w:fill="E9E9E9"/>'.format(nsdecls('w')))
            info_table.rows[i].cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        
        # Thêm dữ liệu vào bảng thông tin
        cells = info_table.rows[0].cells
        cells[0].text = "Họ và tên"
        cells[1].text = student_name
        
        cells = info_table.rows[1].cells
        cells[0].text = "Email"
        cells[1].text = student_email
        
        cells = info_table.rows[2].cells
        cells[0].text = "Lớp"
        cells[1].text = student_class
        
        # Xử lý timestamp tương thích với cả hai kiểu dữ liệu (số và chuỗi ISO)
        submission_time = "Không xác định"
        if isinstance(submission.get("timestamp"), (int, float)):
            try:
                submission_time = datetime.fromtimestamp(submission.get("timestamp")).strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
        else:
            try:
                dt = datetime.fromisoformat(submission.get("timestamp", "").replace("Z", "+00:00"))
                submission_time = dt.strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
        
        cells = info_table.rows[3].cells
        cells[0].text = "Thời gian nộp"
        cells[1].text = submission_time
        
        # Tính toán thông tin về bài làm
        total_correct = 0
        total_questions = len(questions)
        calculated_total_score = 0  # Tổng điểm tính lại từ đầu
        multiple_choice_score = 0  # Điểm trắc nghiệm (Checkbox + Combobox)
        essay_score = 0  # Điểm tự luận (Essay)
        multiple_choice_questions = 0  # Số câu trắc nghiệm
        essay_questions = 0  # Số câu tự luận
        
        doc.add_heading("Chi tiết câu trả lời", level=2)
        
        # Tạo bảng chi tiết câu trả lời - cải thiện layout với cột rộng hợp lý
        # Thêm cột cho loại câu hỏi
        answers_table = doc.add_table(rows=1, cols=5, style='Table Grid')
        
        # Thiết lập độ rộng tương đối cho các cột
        col_widths = [2.5, 2, 2, 1, 0.8]  # Tỷ lệ tương đối
        for i, width in enumerate(col_widths):
            for cell in answers_table.columns[i].cells:
                cell.width = Inches(width)
        
        # Thêm tiêu đề cho bảng với định dạng rõ ràng
        header_cells = answers_table.rows[0].cells
        headers = ["Câu hỏi", "Đáp án của học viên", "Đáp án đúng", "Kết quả", "Điểm"]
        
        # Tạo nền xám cho hàng tiêu đề
        for i, cell in enumerate(header_cells):
            cell.text = headers[i]
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Đảm bảo có runs trước khi truy cập
                if not paragraph.runs:
                    paragraph.add_run(headers[i])
                for run in paragraph.runs:
                    run.bold = True
            # Thêm màu nền
            shading_elm = parse_xml(r'<w:shd {} w:fill="E9E9E9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Đảm bảo responses đúng định dạng - parse từ JSON string nếu cần
        responses = submission.get("responses", {})
        if isinstance(responses, str):
            try:
                responses = json.loads(responses)
            except json.JSONDecodeError as e:
                print(f"Lỗi khi parse responses JSON: {e}")
                responses = {}
        
        # Đảm bảo responses là dict
        if not isinstance(responses, dict):
            print(f"Warning: responses không phải dict, type: {type(responses)}")
            responses = {}
        
        # Validate và normalize questions trước khi xử lý
        normalized_questions = []
        for q in questions:
            # Đảm bảo q là dict và có các trường cần thiết
            if not isinstance(q, dict):
                continue
            
            # Parse answers nếu là string
            if isinstance(q.get("answers"), str):
                try:
                    q["answers"] = json.loads(q["answers"])
                except:
                    try:
                        # Fallback: nếu không phải JSON, giữ nguyên
                        q["answers"] = [q["answers"]]
                    except:
                        q["answers"] = []
            
            # Parse correct nếu là string
            if isinstance(q.get("correct"), str):
                try:
                    q["correct"] = json.loads(q["correct"])
                except:
                    try:
                        # Thử parse dạng "1,2,3"
                        q["correct"] = [int(x.strip()) for x in q["correct"].split(",")]
                    except:
                        q["correct"] = []
            
            # Đảm bảo answers và correct là list
            if not isinstance(q.get("answers"), list):
                q["answers"] = []
            if not isinstance(q.get("correct"), list):
                q["correct"] = []
            
            normalized_questions.append(q)
        
        # Sử dụng questions đã được normalize
        questions = normalized_questions
        
        # Thêm dữ liệu câu trả lời với định dạng cải thiện
        for q in questions:
            q_id = str(q.get("id", ""))
            
            # Đáp án người dùng - lấy từ responses (đã parse từ JSON)
            user_ans = responses.get(q_id, [])
            
            # Đảm bảo user_ans là list
            if not isinstance(user_ans, list):
                if user_ans is not None:
                    user_ans = [user_ans]
                else:
                    user_ans = []
            
            # Kiểm tra đúng/sai sử dụng hàm từ database_helper (không dùng mock)
            # Đảm bảo import đúng hàm từ database_helper
            try:
                from database_helper import check_answer_correctness as db_check_answer
                is_correct = db_check_answer(user_ans, q)
            except ImportError:
                # Fallback nếu không import được - nhưng nên có cảnh báo
                print("Warning: Không thể import check_answer_correctness từ database_helper, sử dụng mock function")
            is_correct = check_answer_correctness(user_ans, q)
            
            # Tính điểm cho câu hỏi này theo từng loại
            q_type = q.get("type", "")
            if is_correct:
                total_correct += 1
                result = "Đúng"
                points = q.get("score", 0)
            else:
                result = "Sai"
                points = 0
            
            # Phân loại điểm: trắc nghiệm hoặc tự luận
            if q_type in ["Checkbox", "Combobox"]:
                # Câu hỏi trắc nghiệm
                multiple_choice_questions += 1
                if is_correct:
                    multiple_choice_score += points
            elif q_type == "Essay":
                # Câu hỏi tự luận
                essay_questions += 1
                if is_correct:
                    essay_score += points
            
            # Cộng vào tổng điểm
            calculated_total_score += points
            
            # Thêm hàng mới vào bảng
            row_cells = answers_table.add_row().cells
            
            # Thêm thông tin câu hỏi
            row_cells[0].text = f"Câu {q.get('id', '')}: {q.get('question', '')}"
            
            # Xử lý nội dung đáp án dựa trên loại câu hỏi
            if q.get("type") == "Essay":
                # Đối với câu hỏi tự luận
                essay_answer = user_ans[0] if user_ans else "Không trả lời"
                row_cells[1].text = essay_answer
                
                # Đối với câu hỏi tự luận, không có đáp án đúng
                row_cells[2].text = "Câu hỏi tự luận"
                
                # Kết quả dựa trên việc học viên có trả lời hay không
                result = "Đã trả lời" if is_correct else "Không trả lời"
                row_cells[3].text = result
            else:
                # Đối với câu hỏi trắc nghiệm
                row_cells[1].text = ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời"
                
                # Chuẩn bị đáp án đúng
                q_correct = q.get("correct", [])
                q_answers = q.get("answers", [])
                
                if isinstance(q_correct, str):
                    try:
                        q_correct = json.loads(q_correct)
                    except:
                        try:
                            q_correct = [int(x.strip()) for x in q_correct.split(",")]
                        except:
                            q_correct = []
                
                if isinstance(q_answers, str):
                    try:
                        q_answers = json.loads(q_answers)
                    except:
                        q_answers = [q_answers]
                
                try:
                    expected = [q_answers[i - 1] for i in q_correct]
                except (IndexError, TypeError):
                    expected = ["Lỗi đáp án"]
                
                row_cells[2].text = ", ".join([str(a) for a in expected])
                row_cells[3].text = result
            
            # Đặt màu cho kết quả
            for paragraph in row_cells[3].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not paragraph.runs:
                    paragraph.add_run(result)
                run = paragraph.runs[0]
                if is_correct:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Màu xanh lá cho đúng
                    run.bold = True
                else:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Màu đỏ cho sai
                    run.bold = True
            
            row_cells[4].text = str(points)
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tính điểm tối đa cho từng loại
        max_multiple_choice = sum([q.get("score", 0) for q in questions if q.get("type") in ["Checkbox", "Combobox"]])
        max_essay = sum([q.get("score", 0) for q in questions if q.get("type") == "Essay"])
        
        # Thêm tổng kết với định dạng rõ ràng - bao gồm điểm trắc nghiệm và tự luận
        doc.add_heading("Tổng kết", level=2)
        summary_table = doc.add_table(rows=6, cols=2, style='Table Grid')
        
        # Thiết lập độ rộng cho bảng tổng kết
        for cell in summary_table.columns[0].cells:
            cell.width = Inches(1.5)
        for cell in summary_table.columns[1].cells:
            cell.width = Inches(3.0)
        
        # Thêm màu nền cho cột tiêu đề
        for i in range(6):
            cell = summary_table.rows[i].cells[0]
            paragraph = cell.paragraphs[0]
            if not paragraph.runs:
                paragraph.add_run(cell.text if cell.text else '')
            paragraph.runs[0].font.bold = True
            shading_elm = parse_xml(r'<w:shd {} w:fill="E9E9E9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Thông tin tổng quan
        cells = summary_table.rows[0].cells
        cells[0].text = "Số câu đúng"
        cells[1].text = f"{total_correct}/{total_questions}"
        
        cells = summary_table.rows[1].cells
        cells[0].text = "Điểm trắc nghiệm"
        cells[1].text = f"{multiple_choice_score}/{max_multiple_choice}" if max_multiple_choice > 0 else "0/0"
        
        cells = summary_table.rows[2].cells
        cells[0].text = "Điểm tự luận"
        cells[1].text = f"{essay_score}/{max_essay}" if max_essay > 0 else "0/0"
        
        cells = summary_table.rows[3].cells
        cells[0].text = "Tổng điểm"
        # Sử dụng điểm đã tính lại để đảm bảo chính xác
        cells[1].text = f"{calculated_total_score}/{max_possible}"
        
        cells = summary_table.rows[4].cells
        cells[0].text = "Tỷ lệ đúng"
        cells[1].text = f"{(total_correct/total_questions*100):.1f}%" if total_questions > 0 else "0%"
        
        cells = summary_table.rows[5].cells
        cells[0].text = "Tỷ lệ điểm"
        cells[1].text = f"{(calculated_total_score/max_possible*100):.1f}%" if max_possible > 0 else "0%"
        
        # Thêm chân trang
        doc.add_paragraph()
        footer = doc.add_paragraph("Xuất báo cáo từ Hệ thống Khảo sát & Đánh giá")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_footer = doc.add_paragraph(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        time_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Lưu tệp vào buffer - đảm bảo cách xử lý đúng
        temp_buffer = io.BytesIO()
        try:
            # Lưu document vào buffer tạm
            doc.save(temp_buffer)
            temp_buffer.flush()
            temp_buffer.seek(0)
            
            # Tạo buffer chính và copy dữ liệu
            buffer = io.BytesIO()
            buffer.write(temp_buffer.getvalue())
            temp_buffer.close()
            
            # Đảm bảo buffer ở đầu
            buffer.seek(0)
            
            # Kiểm tra nội dung
            content = buffer.getvalue()
            if not content or len(content) < 100:
                raise ValueError(f"DOCX buffer is empty or too small (length: {len(content) if content else 0})")
            
            # Kiểm tra signature DOCX (PK = ZIP format)
            if content[:2] != b'PK':
                raise ValueError("DOCX buffer does not contain valid DOCX file (missing ZIP signature)")
            
            # Đảm bảo buffer ở đầu để sẵn sàng đọc
            buffer.seek(0)
            return buffer
        except Exception as save_error:
            print(f"Lỗi khi lưu DOCX vào buffer: {save_error}")
            import traceback
            traceback.print_exc()
            try:
                temp_buffer.close()
            except:
                pass
            raise
            
    except Exception as e:
        print(f"Lỗi khi tạo báo cáo DOCX: {str(e)}")
        import traceback
        traceback.print_exc()
        st.error(f"Không thể tạo báo cáo DOCX: {str(e)}")
        # Trả về None để dễ kiểm tra lỗi
        return None



def display_overview_tab(submissions=None, students=None, questions=None, max_possible=0):
    """Hiển thị tab tổng quan"""
    if submissions is None:
        submissions = []
    if students is None:
        students = []
    if questions is None:
        questions = []
        
    st.subheader("Tổng quan kết quả")
    

def create_student_report_pdf_fpdf(student_name, student_email, student_class, submission, questions, max_possible):
    """Tạo báo cáo chi tiết bài làm của học viên dạng PDF sử dụng FPDF2 với hỗ trợ Unicode"""
    # Buffer sẽ được tạo khi lưu PDF, không tạo ở đây
    try:
        # Tạo PDF mới với hỗ trợ Unicode
        title = f"Báo cáo chi tiết - {student_name}"
        pdf = create_unicode_pdf(title=title)
        
        if pdf is None:
            raise Exception("Không thể tạo đối tượng PDF")
        
        pdf.add_page()
        
        # Lấy tên font đã được add (nếu có)
        font_name = getattr(pdf, '_active_font_name', 'DejaVu')
        
        # Thiết lập font cho tiêu đề - sử dụng font đã được add
        try:
            pdf.set_font(font_name, 'B', 16)
        except:
            try:
                pdf.set_font('Arial', 'B', 16)  # Fallback to built-in
                font_name = 'Arial'
            except:
                pdf.set_font('Helvetica', 'B', 16)  # Last resort
                font_name = 'Helvetica'
        
        pdf.cell(0, 10, title, 0, 1, 'C')
        
        # Thêm thời gian báo cáo
        try:
            pdf.set_font(font_name, 'I', 10)
        except:
            pdf.set_font(font_name, '', 10)  # Nếu không có italic, dùng regular
            
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        pdf.cell(0, 5, f'Thoi gian xuat bao cao: {timestamp}', 0, 1, 'R')
        pdf.ln(5)
        
        # Tính toán thông tin về bài làm
        total_correct = 0
        total_questions = len(questions)
        calculated_total_score = 0  # Tổng điểm tính lại từ đầu
        multiple_choice_score = 0  # Điểm trắc nghiệm (Checkbox + Combobox)
        essay_score = 0  # Điểm tự luận (Essay)
        multiple_choice_questions = 0  # Số câu trắc nghiệm
        essay_questions = 0  # Số câu tự luận
        
        # Đảm bảo responses đúng định dạng - parse từ JSON string nếu cần
        responses = submission.get("responses", {})
        if isinstance(responses, str):
            try:
                responses = json.loads(responses)
            except json.JSONDecodeError as e:
                print(f"Lỗi khi parse responses JSON trong PDF: {e}")
                responses = {}
        
        # Đảm bảo responses là dict
        if not isinstance(responses, dict):
            print(f"Warning: responses không phải dict trong PDF, type: {type(responses)}")
            responses = {}
        
        # Validate và normalize questions trước khi xử lý
        normalized_questions = []
        for q in questions:
            # Đảm bảo q là dict và có các trường cần thiết
            if not isinstance(q, dict):
                continue
            
            # Parse answers nếu là string
            if isinstance(q.get("answers"), str):
                try:
                    q["answers"] = json.loads(q["answers"])
                except:
                    try:
                        q["answers"] = [q["answers"]]
                    except:
                        q["answers"] = []
            
            # Parse correct nếu là string
            if isinstance(q.get("correct"), str):
                try:
                    q["correct"] = json.loads(q["correct"])
                except:
                    try:
                        q["correct"] = [int(x.strip()) for x in q["correct"].split(",")]
                    except:
                        q["correct"] = []
            
            # Đảm bảo answers và correct là list
            if not isinstance(q.get("answers"), list):
                q["answers"] = []
            if not isinstance(q.get("correct"), list):
                q["correct"] = []
            
            normalized_questions.append(q)
        
        # Sử dụng questions đã được normalize
        questions = normalized_questions
        
        # Xử lý timestamp
        submission_time = "Không xác định"
        if isinstance(submission.get("timestamp"), (int, float)):
            try:
                submission_time = datetime.fromtimestamp(submission.get("timestamp")).strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
        else:
            try:
                dt = datetime.fromisoformat(submission.get("timestamp", "").replace("Z", "+00:00"))
                submission_time = dt.strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
        
        # Thông tin học viên
        _set_font_safe(pdf, 'B', 12)
        pdf.cell(0, 10, 'Thong tin hoc vien', 0, 1, 'L')
        
        # Bảng thông tin học viên
        _set_font_safe(pdf, '', 10)
        info_width = 190
        col1_width = 50
        col2_width = info_width - col1_width
        
        # Tạo khung thông tin học viên - convert sang ASCII-safe nếu cần
        pdf.set_fill_color(240, 240, 240)
        font_name = getattr(pdf, '_active_font_name', 'DejaVu')
        
        # Helper để convert text sang ASCII-safe nếu font không hỗ trợ Unicode tốt
        def safe_text(text, max_len=100):
            text_str = str(text)
            if len(text_str) > max_len:
                text_str = text_str[:max_len-3] + "..."
            if font_name not in ['DejaVu', 'Arial']:
                # Convert sang ASCII-safe cho built-in fonts
                text_str = ''.join(c if ord(c) < 128 else '?' for c in text_str)
            return text_str
        
        pdf.cell(col1_width, 10, 'Ho va ten', 1, 0, 'L', 1)
        pdf.cell(col2_width, 10, safe_text(student_name), 1, 1, 'L')
        
        pdf.cell(col1_width, 10, 'Email', 1, 0, 'L', 1)
        pdf.cell(col2_width, 10, safe_text(student_email), 1, 1, 'L')
        
        pdf.cell(col1_width, 10, 'Lop', 1, 0, 'L', 1)
        pdf.cell(col2_width, 10, safe_text(student_class), 1, 1, 'L')
        
        pdf.cell(col1_width, 10, 'Thoi gian nop', 1, 0, 'L', 1)
        pdf.cell(col2_width, 10, safe_text(submission_time), 1, 1, 'L')
        
        pdf.ln(5)
        
        # Chi tiết câu trả lời
        _set_font_safe(pdf, 'B', 12)
        pdf.cell(0, 10, 'Chi tiet cau tra loi', 0, 1, 'L')
        
        # Tiêu đề bảng chi tiết
        _set_font_safe(pdf, 'B', 9)
        pdf.set_fill_color(240, 240, 240)
        
        # Xác định độ rộng cột - điều chỉnh phù hợp với nội dung
        q_width = 80
        user_width = 35
        correct_width = 35
        result_width = 20
        points_width = 20
        
        # Kiểm tra tổng độ rộng cột
        total_width = q_width + user_width + correct_width + result_width + points_width
        page_width = 210 - 20  # A4 width - margin
        
        # Điều chỉnh nếu vượt quá chiều rộng trang
        if total_width > page_width:
            scale = page_width / total_width
            q_width *= scale
            user_width *= scale
            correct_width *= scale
            result_width *= scale
            points_width *= scale

        # Đảm bảo mỗi cột tối thiểu đủ để render 1 ký tự
        min_char = pdf.get_string_width('W') + 4
        q_width = max(q_width, min_char)
        user_width = max(user_width, min_char)
        correct_width = max(correct_width, min_char)
        result_width = max(result_width, min_char)
        points_width = max(points_width, min_char)
        
        # Vẽ header bảng
        pdf.cell(q_width, 10, 'Cau hoi', 1, 0, 'C', 1)
        pdf.cell(user_width, 10, 'Dap an HV', 1, 0, 'C', 1)
        pdf.cell(correct_width, 10, 'Dap an dung', 1, 0, 'C', 1)
        pdf.cell(result_width, 10, 'Ket qua', 1, 0, 'C', 1)
        pdf.cell(points_width, 10, 'Diem', 1, 1, 'C', 1)
        
        # Vẽ dữ liệu câu trả lời
        _set_font_safe(pdf, '', 9)
        
        for q in questions:
            q_id = str(q.get("id", ""))
            
            # Đáp án người dùng - lấy từ responses (đã parse từ JSON)
            user_ans = responses.get(q_id, [])
            
            # Đảm bảo user_ans là list
            if not isinstance(user_ans, list):
                if user_ans is not None:
                    user_ans = [user_ans]
                else:
                    user_ans = []
            
            # Kiểm tra đúng/sai sử dụng hàm từ database_helper (không dùng mock)
            try:
                from database_helper import check_answer_correctness as db_check_answer
                is_correct = db_check_answer(user_ans, q)
            except ImportError:
                # Fallback nếu không import được - nhưng nên có cảnh báo
                print("Warning: Không thể import check_answer_correctness từ database_helper, sử dụng mock function")
            is_correct = check_answer_correctness(user_ans, q)
            
            # Tính điểm cho câu hỏi này theo từng loại
            q_type = q.get("type", "")
            if is_correct:
                total_correct += 1
                points = q.get("score", 0)
            else:
                points = 0
            
            # Phân loại điểm: trắc nghiệm hoặc tự luận
            if q_type in ["Checkbox", "Combobox"]:
                # Câu hỏi trắc nghiệm
                multiple_choice_questions += 1
                if is_correct:
                    multiple_choice_score += points
            elif q_type == "Essay":
                # Câu hỏi tự luận
                essay_questions += 1
                if is_correct:
                    essay_score += points
            
            # Cộng vào tổng điểm
            calculated_total_score += points
            
            # Chuẩn bị nội dung dựa trên loại câu hỏi
            # Helper để convert text sang ASCII-safe
            font_name = getattr(pdf, '_active_font_name', 'DejaVu')
            def safe_text_pdf(text, max_len=200):
                text_str = str(text)
                if len(text_str) > max_len:
                    text_str = text_str[:max_len-3] + "..."
                if font_name not in ['DejaVu', 'Arial']:
                    # Convert sang ASCII-safe cho built-in fonts
                    text_str = ''.join(c if ord(c) < 128 else '?' for c in text_str)
                return text_str
            
            question_text = safe_text_pdf(f"Câu {q.get('id', '')}: {q.get('question', '')}", 150)
            
            if q.get("type") == "Essay":
                # Đối với câu hỏi tự luận
                essay_answer = user_ans[0] if user_ans else "Không trả lời"
                user_answer_text = safe_text_pdf(essay_answer)
                correct_answer_text = "Cau hoi tu luan"
                result = "Da tra loi" if is_correct else "Khong tra loi"
            else:
                # Đối với câu hỏi trắc nghiệm
                user_answer_text = safe_text_pdf(", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời")
                
                # Chuẩn bị đáp án đúng
                q_correct = q.get("correct", [])
                q_answers = q.get("answers", [])
                
                if isinstance(q_correct, str):
                    try:
                        q_correct = json.loads(q_correct)
                    except:
                        try:
                            q_correct = [int(x.strip()) for x in q_correct.split(",")]
                        except:
                            q_correct = []
                
                if isinstance(q_answers, str):
                    try:
                        q_answers = json.loads(q_answers)
                    except:
                        q_answers = [q_answers]
                
                try:
                    expected = [q_answers[i - 1] for i in q_correct]
                except (IndexError, TypeError):
                    expected = ["Loi dap an"]
                
                correct_answer_text = safe_text_pdf(", ".join([str(a) for a in expected]) if expected else "Khong co dap an")
                result = "Dung" if is_correct else "Sai"
            
            # Tính chiều cao cần thiết cho từng ô bằng split_only
            line_h = 5
            q_height = _measure_multicell_height(pdf, q_width, question_text, line_height=line_h)
            user_height = _measure_multicell_height(pdf, user_width, user_answer_text, line_height=line_h)
            correct_height = _measure_multicell_height(pdf, correct_width, correct_answer_text, line_height=line_h)
            row_height = max(7, q_height, user_height, correct_height)
            
            # Lưu vị trí x hiện tại
            x = pdf.get_x()
            y = pdf.get_y()
            
            # Kiểm tra nếu chiều cao của dòng này sẽ vượt quá trang
            if y + row_height > pdf.page_break_trigger:
                pdf.add_page()
                y = pdf.get_y()
            
            # Vẽ câu hỏi (multi_cell để tự động xuống dòng)
            pdf.set_text_color(0, 0, 0)
            pdf.set_xy(x, y)
            pdf.multi_cell(q_width, line_h, question_text, border=1, align='L')
            
            # Ghi nhớ y sau cột 1 để canh hàng
            y_after_q = pdf.get_y()
            
            # Vẽ đáp án của học viên
            pdf.set_xy(x + q_width, y)
            pdf.multi_cell(user_width, line_h, user_answer_text, border=1, align='L')
            
            # Vẽ đáp án đúng
            pdf.set_xy(x + q_width + user_width, y)
            pdf.multi_cell(correct_width, line_h, correct_answer_text, border=1, align='L')
            
            # Vẽ kết quả với màu tương ứng
            pdf.set_xy(x + q_width + user_width + correct_width, y)
            if is_correct:
                pdf.set_text_color(0, 128, 0)  # Màu xanh lá
            else:
                pdf.set_text_color(255, 0, 0)  # Màu đỏ
            pdf.cell(result_width, row_height, result, 1, 0, 'C')
            
            # Vẽ điểm
            pdf.set_text_color(0, 0, 0)
            pdf.set_xy(x + q_width + user_width + correct_width + result_width, y)
            pdf.cell(points_width, row_height, str(points), 1, 1, 'C')
        
        pdf.ln(5)
        
        # Tính điểm tối đa cho từng loại
        max_multiple_choice = sum([q.get("score", 0) for q in questions if q.get("type") in ["Checkbox", "Combobox"]])
        max_essay = sum([q.get("score", 0) for q in questions if q.get("type") == "Essay"])
        
        # Tổng kết - bao gồm điểm trắc nghiệm và tự luận
        _set_font_safe(pdf, 'B', 12)
        pdf.cell(0, 10, 'Tong ket', 0, 1, 'L')
        
        # Bảng tổng kết
        _set_font_safe(pdf, '', 10)
        pdf.set_fill_color(240, 240, 240)
        
        summary_col1 = 50
        summary_col2 = 140
        
        pdf.cell(summary_col1, 10, 'So cau dung', 1, 0, 'L', 1)
        pdf.cell(summary_col2, 10, f"{total_correct}/{total_questions}", 1, 1, 'L')
        
        pdf.cell(summary_col1, 10, 'Diem trac nghiem', 1, 0, 'L', 1)
        pdf.cell(summary_col2, 10, f"{multiple_choice_score}/{max_multiple_choice}" if max_multiple_choice > 0 else "0/0", 1, 1, 'L')
        
        pdf.cell(summary_col1, 10, 'Diem tu luan', 1, 0, 'L', 1)
        pdf.cell(summary_col2, 10, f"{essay_score}/{max_essay}" if max_essay > 0 else "0/0", 1, 1, 'L')
        
        pdf.cell(summary_col1, 10, 'Tong diem', 1, 0, 'L', 1)
        # Sử dụng điểm đã tính lại để đảm bảo chính xác
        pdf.cell(summary_col2, 10, f"{calculated_total_score}/{max_possible}", 1, 1, 'L')
        
        pdf.cell(summary_col1, 10, 'Ty le dung', 1, 0, 'L', 1)
        percent = total_correct/total_questions*100 if total_questions > 0 else 0
        pdf.cell(summary_col2, 10, f"{percent:.1f}% {'(Dat)' if percent >= 50 else '(Chua dat)'}", 1, 1, 'L')
        
        pdf.cell(summary_col1, 10, 'Ty le diem', 1, 0, 'L', 1)
        score_percent = calculated_total_score/max_possible*100 if max_possible > 0 else 0
        pdf.cell(summary_col2, 10, f"{score_percent:.1f}%", 1, 1, 'L')
        
        # Lưu PDF vào buffer - luôn sử dụng dest='S' để lấy bytes
        buffer = io.BytesIO()  # Đảm bảo buffer sạch
        try:
            # Sử dụng dest='S' để lấy bytes trực tiếp (cách an toàn nhất)
            pdf_bytes = pdf.output(dest='S')
            
            # Xử lý các kiểu dữ liệu khác nhau
            if isinstance(pdf_bytes, bytes):
                buffer.write(pdf_bytes)
            elif isinstance(pdf_bytes, str):
                # Nếu là string, encode sang bytes
                buffer.write(pdf_bytes.encode('latin-1', errors='ignore'))
            elif hasattr(pdf_bytes, '__iter__'):
                # Nếu là iterable, chuyển thành bytes
                buffer.write(bytes(pdf_bytes))
            else:
                # Fallback: thử encode sang string rồi bytes
                buffer.write(str(pdf_bytes).encode('latin-1', errors='ignore'))
            
            buffer.flush()
            
        except Exception as output_error:
            print(f"Lỗi khi output PDF học viên: {output_error}")
            # Thử lại với dest=buffer
            try:
                buffer = io.BytesIO()
                pdf.output(dest=buffer)
                buffer.flush()
            except Exception:
                raise ValueError(f"Không thể tạo PDF output: {str(output_error)}")
        
        # Đảm bảo buffer ở đầu và kiểm tra dữ liệu
        buffer.seek(0)
        content = buffer.getvalue()
        
        # Kiểm tra signature PDF
        if not content or len(content) < 100:
            # Thử đọc lại bằng read()
            buffer.seek(0)
            content = buffer.read()
            buffer.seek(0)
            
            if not content or len(content) < 100:
                raise ValueError(f"PDF buffer is empty or too small (length: {len(content) if content else 0}). Có thể PDF chưa được tạo đúng cách.")
        
        if not content.startswith(b'%PDF'):
            raise ValueError(f"PDF buffer does not contain valid PDF (missing %PDF signature). Buffer đầu: {content[:50] if content else 'None'}")
        
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        print(f"Lỗi khi tạo báo cáo PDF học viên: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Tạo báo cáo đơn giản nếu gặp lỗi
        try:
            buffer = io.BytesIO()  # Tạo buffer mới
            simple_pdf = FPDF()
            simple_pdf.add_page()
            simple_pdf.set_font('Arial', 'B', 16)
            simple_pdf.cell(0, 10, f'Bao cao chi tiet - {student_name}', 0, 1, 'C')
            simple_pdf.set_font('Arial', '', 10)
            error_msg = f'Khong the hien thi bao cao chi tiet. Vui long su dung dinh dang DOCX hoac Excel.\nLoi: {str(e)[:100]}'
            simple_pdf.multi_cell(0, 10, error_msg, 0, 'L')
            
            try:
                pdf_bytes = simple_pdf.output(dest='S')
                if isinstance(pdf_bytes, str):
                    pdf_bytes = pdf_bytes.encode('latin-1')
                buffer.write(pdf_bytes)
                buffer.flush()
            except Exception:
                simple_pdf.output(dest=buffer)
                buffer.flush()
            
            buffer.seek(0)
            
            # Kiểm tra lại
            content = buffer.getvalue()
            if content and len(content) > 100 and content.startswith(b'%PDF'):
                return buffer
            else:
                return None
        except Exception as e2:
            print(f"Không thể tạo báo cáo thay thế: {str(e2)}")
            return None

def display_overview_tab(submissions=None, students=None, questions=None, max_possible=0):
    """Hiển thị tab tổng quan"""
    if submissions is None:
        submissions = []
    if students is None:
        students = []
    if questions is None:
        questions = []
        
    st.subheader("Tổng quan kết quả")
    
    # Thống kê cơ bản
    total_submissions = len(submissions)
    if total_submissions > 0:
        scores = [s.get("score", 0) for s in submissions]
        avg_score = sum(scores) / total_submissions
        max_score = max(scores) if scores else 0
    else:
        avg_score = 0
        max_score = 0
        
    total_users = len(set([s.get("user_email") for s in submissions])) if submissions else 0
    
    # Hiển thị metrics
    col1, col2, col3 = st.columns(3)
    col1.metric("📝 Tổng số bài nộp", total_submissions)
    col1.metric("👥 Số học viên đã làm", total_users)
    
    col2.metric("📊 Điểm trung bình", f"{avg_score:.2f}/{max_possible}")
    col2.metric("🏆 Điểm cao nhất", f"{max_score}/{max_possible}")
    
    col3.metric("📋 Số câu hỏi", len(questions))
    col3.metric("👨‍🎓 Tổng số học viên", len(students))
    
    # Biểu đồ điểm số theo thời gian
    st.subheader("Điểm số theo thời gian")
    
    # Chuẩn bị dữ liệu
    time_data = []
    for s in submissions:
        # Xử lý timestamp
        try:
            if isinstance(s.get("timestamp"), (int, float)):
                submit_time = datetime.fromtimestamp(s.get("timestamp"))
            else:
                submit_time = datetime.fromisoformat(s.get("timestamp", "").replace("Z", "+00:00"))
            
            time_data.append({
                "timestamp": submit_time,
                "score": s.get("score", 0),
                "user": s.get("user_email", "")
            })
        except:
            # Bỏ qua bài nộp có timestamp không hợp lệ
            pass
    
    if time_data:
        df_time = pd.DataFrame(time_data)
        if not df_time.empty:
            df_time = df_time.sort_values("timestamp")
            
            # Vẽ biểu đồ
            fig, ax = plt.subplots(figsize=(10, 5))
            ax.plot(df_time["timestamp"], df_time["score"], marker='o')
            ax.set_ylabel("Điểm số")
            ax.set_xlabel("Thời gian nộp bài")
            ax.grid(True, linestyle='--', alpha=0.7)
            
            # Giảm số lượng tick trên trục x
            max_ticks = 6
            if len(df_time) > max_ticks:
                stride = len(df_time) // max_ticks
                plt.xticks(df_time["timestamp"][::stride])
            
            # Sử dụng constrained_layout thay vì tight_layout
            fig.set_constrained_layout(True)
            st.pyplot(fig)
            
            # Hiển thị phân phối điểm
            st.subheader("Phân phối điểm số")
            if submissions:
                scores = [s.get("score", 0) for s in submissions]
                fig, ax = plt.subplots(figsize=(10, 5))
                ax.hist(scores, bins=min(10, len(set(scores))), alpha=0.7, color='skyblue', edgecolor='black')
                ax.set_xlabel("Điểm số")
                ax.set_ylabel("Số lượng bài nộp")
                ax.grid(True, linestyle='--', alpha=0.3)
                fig.set_constrained_layout(True)
                st.pyplot(fig)
    else:
        st.info("Không có đủ dữ liệu để vẽ biểu đồ theo thời gian.")

def display_student_tab(submissions=None, students=None, questions=None, max_possible=0):
    """Hiển thị tab theo học viên"""
    if submissions is None:
        submissions = []
    if students is None:
        students = []
    if questions is None:
        questions = []
        
    # Đảm bảo load lại students nếu chưa có - bao gồm tất cả roles
    if not students:
        try:
            # Thử dùng hàm get_all_students nếu có
            try:
                from database_helper import get_all_students
                students = get_all_students()
            except ImportError:
                # Fallback: load tất cả users với các role
                students = get_all_users(role=["Học viên", "student", "admin"])
            if not students:
                # Fallback cuối: load tất cả users và filter
                all_users = get_all_users(role=None)
                if all_users:
                    valid_roles = ["Học viên", "student", "admin"]
                    students = [u for u in all_users if u.get("role") in valid_roles]
        except Exception as e:
            st.error(f"❌ Lỗi khi load danh sách users: {str(e)}")
            students = []
    
    st.subheader("Chi tiết theo học viên")
    
    # Tạo dict để lookup nhanh hơn (email -> student info)
    students_dict = {student.get("email", ""): student for student in students if student.get("email")}
    
    if not students_dict and students:
        st.warning("⚠️ Không thể tạo danh sách học viên. Kiểm tra dữ liệu email.")
    
    # Tạo DataFrame từ dữ liệu
    user_data = []
    for s in submissions:
        try:
            user_email = s.get("user_email", "")
            # Tìm thông tin học viên từ dict (nhanh hơn)
            student_info = students_dict.get(user_email)
            
            if student_info:
                full_name = student_info.get("full_name", "Không xác định")
                class_name = student_info.get("class", "Chưa phân lớp")
            else:
                # Nếu không tìm thấy trong students list, thử query trực tiếp từ DB
                try:
                    from database_helper import get_all_students, get_all_users
                    # Thử dùng get_all_students trước
                    try:
                        all_students = get_all_students()
                    except:
                        all_students = get_all_users(role=["Học viên", "student", "admin"])
                    student_info = next((st for st in all_students if st.get("email") == user_email), None)
                    if student_info:
                        full_name = student_info.get("full_name", "Không xác định")
                        class_name = student_info.get("class", "Chưa phân lớp")
                        # Cập nhật vào dict để dùng lần sau
                        students_dict[user_email] = student_info
                    else:
                        full_name = "Không xác định"
                        class_name = "Không xác định"
                except:
                    full_name = "Không xác định"
                    class_name = "Không xác định"
            
            # Xử lý timestamp
            submission_time = "Không xác định"
            try:
                if isinstance(s.get("timestamp"), (int, float)):
                    submission_time = datetime.fromtimestamp(s.get("timestamp")).strftime("%H:%M:%S %d/%m/%Y")
                else:
                    dt = datetime.fromisoformat(s.get("timestamp", "").replace("Z", "+00:00"))
                    submission_time = dt.strftime("%H:%M:%S %d/%m/%Y")
            except:
                pass
            
            user_data.append({
                "email": s.get("user_email", ""),
                "full_name": full_name,
                "class": class_name,
                "submission_id": s.get("id", ""),
                "timestamp": submission_time,
                "score": s.get("score", 0),
                "max_score": max_possible,
                "percent": f"{(s.get('score', 0)/max_possible*100):.1f}%" if max_possible > 0 else "N/A"
            })
        except Exception as e:
            st.error(f"Lỗi khi xử lý dữ liệu học viên: {str(e)}")
    
    if user_data:
        df_users = pd.DataFrame(user_data)
        
        # Lọc theo email hoặc lớp
        col1, col2 = st.columns(2)
        with col1:
            user_filter = st.selectbox(
                "Chọn học viên để xem chi tiết:",
                options=["Tất cả"] + sorted(list(set([u.get("email", "") for u in user_data]))),
                key="user_filter_tab2"
            )
        
        with col2:
            unique_classes = [u.get("class", "") for u in user_data if u.get("class") != "Không xác định"]
            class_filter = st.selectbox(
                "Lọc theo lớp:",
                options=["Tất cả"] + sorted(list(set(unique_classes))),
                key="class_filter_tab2"
            )
        
        # Áp dụng bộ lọc
        df_filtered = df_users
        
        if user_filter != "Tất cả":
            df_filtered = df_filtered[df_filtered["email"] == user_filter]
        
        if class_filter != "Tất cả":
            df_filtered = df_filtered[df_filtered["class"] == class_filter]
        
        # Hiển thị bảng với đầy đủ cột
        columns_to_show = ["email", "full_name", "class", "timestamp", "score", "max_score", "percent"]
        if "submission_id" in df_filtered.columns:
            columns_to_show.insert(0, "submission_id")
        
        # Đảm bảo tất cả cột tồn tại
        available_columns = [col for col in columns_to_show if col in df_filtered.columns]
        df_display = df_filtered[available_columns].copy()
        
        # Đổi tên cột cho dễ đọc
        column_mapping = {
            "email": "Email",
            "full_name": "Họ và tên",
            "class": "Lớp",
            "timestamp": "Thời gian nộp",
            "score": "Điểm số",
            "max_score": "Điểm tối đa",
            "percent": "Tỷ lệ",
            "submission_id": "ID bài nộp"
        }
        df_display = df_display.rename(columns=column_mapping)
        
        st.dataframe(
            df_display.sort_values(by="Thời gian nộp", ascending=False),
            width='stretch',
            hide_index=True
        )
        
        # Xem chi tiết một bài nộp cụ thể
        if user_filter != "Tất cả":
            submission_ids = df_filtered["submission_id"].tolist()
            if submission_ids:
                selected_submission = st.selectbox(
                    "Chọn bài nộp để xem chi tiết:",
                    options=submission_ids,
                    key="submission_id_select"
                )
                
                # Tìm bài nộp được chọn
                submission = next((s for s in submissions if str(s.get("id", "")) == str(selected_submission)), None)
                if submission:
                    st.subheader(f"Chi tiết bài nộp #{selected_submission}")
                    
                    total_correct = 0
                    total_questions = len(questions)
                    student_detail_data = []
                    
                    # Đảm bảo responses đúng định dạng
                    responses = submission.get("responses", {})
                    if isinstance(responses, str):
                        try:
                            responses = json.loads(responses)
                        except:
                            responses = {}
                    
                    # Hiển thị câu trả lời chi tiết
                    for q in questions:
                        q_id = str(q.get("id", ""))
                        st.write(f"**Câu {q.get('id', '')}: {q.get('question', '')}**")
                        
                        # Đáp án người dùng - lấy từ responses (đã parse từ JSON)
                        user_ans = responses.get(q_id, [])
                        
                        # Đảm bảo user_ans là list
                        if not isinstance(user_ans, list):
                            if user_ans is not None:
                                user_ans = [user_ans]
                            else:
                                user_ans = []
                        
                        # Chuẩn bị dữ liệu đáp án đúng - parse JSON nếu cần
                        q_correct = q.get("correct", [])
                        q_answers = q.get("answers", [])
                        
                        if isinstance(q_correct, str):
                            try:
                                q_correct = json.loads(q_correct)
                            except:
                                try:
                                    q_correct = [int(x.strip()) for x in q_correct.split(",")]
                                except:
                                    q_correct = []
                        
                        if isinstance(q_answers, str):
                            try:
                                q_answers = json.loads(q_answers)
                            except:
                                q_answers = [q_answers]
                        
                        try:
                            expected = [q_answers[i - 1] for i in q_correct]
                        except (IndexError, TypeError):
                            expected = ["Lỗi đáp án"]
                        
                        # Kiểm tra đúng/sai - sử dụng hàm từ database_helper (không dùng mock)
                        try:
                            from database_helper import check_answer_correctness as db_check_answer
                            is_correct = db_check_answer(user_ans, q)
                        except ImportError:
                        is_correct = check_answer_correctness(user_ans, q)
                        if is_correct:
                            total_correct += 1
                        
                        # Thu thập dữ liệu chi tiết
                        student_detail_data.append({
                            "Câu hỏi": f"Câu {q.get('id', '')}: {q.get('question', '')}",
                            "Đáp án của học viên": ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời",
                            "Đáp án đúng": ", ".join([str(a) for a in expected]),
                            "Kết quả": "Đúng" if is_correct else "Sai",
                            "Điểm": q.get("score", 0) if is_correct else 0
                        })
                        
                        # Hiển thị đáp án của người dùng
                        st.write("Đáp án của học viên:")
                        if not user_ans:
                            st.write("- Không trả lời")
                        else:
                            for ans in user_ans:
                                st.write(f"- {ans}")
                        
                        # Hiển thị kết quả
                        if is_correct:
                            st.success(f"✅ Đúng (+{q.get('score', 0)} điểm)")
                        else:
                            st.error("❌ Sai (0 điểm)")
                            st.write("Đáp án đúng:")
                            for ans in expected:
                                st.write(f"- {ans}")
                        
                        st.divider()
                    
                    # Hiển thị thống kê tổng hợp
                    st.subheader("Tổng kết")
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Số câu đúng", f"{total_correct}/{total_questions}")
                    col2.metric("Điểm số", f"{submission.get('score', 0)}/{max_possible}")
                    col3.metric("Tỷ lệ đúng", f"{(total_correct/total_questions*100):.1f}%" if total_questions > 0 else "0%")
                    
                    # Tạo DataFrame chi tiết
                    df_student_detail = pd.DataFrame(student_detail_data)
                    
                    # Xuất báo cáo chi tiết
                    st.write("### Xuất báo cáo chi tiết")
                    
                    # Người dùng và thông tin
                    student_info = next((student for student in students if student.get("email") == submission.get("user_email")), None)
                    student_name = student_info.get("full_name", "Không xác định") if student_info else "Không xác định"
                    student_class = student_info.get("class", "Không xác định") if student_info else "Không xác định"
                    
                    # Tạo báo cáo chi tiết
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        # Tạo báo cáo dạng DOCX
                        try:
                            docx_buffer = create_student_report_docx(
                                student_name,
                                submission.get("user_email", ""),
                                student_class,
                                submission,
                                questions,
                                max_possible
                            )
                            
                            if docx_buffer is not None:
                                get_download_link_docx(
                                    docx_buffer, 
                                                    f"bao_cao_{student_name.replace(' ', '_')}_{submission.get('id', '')}.docx", 
                                    "📥 Tải xuống báo cáo chi tiết (DOCX)"
                            )
                        except Exception as e:
                            st.error(f"Không thể tạo báo cáo DOCX: {str(e)}")
                            import traceback
                            traceback.print_exc()
                    
                    with col2:
                        # Tạo báo cáo dạng PDF
                        try:
                            try:
                                pdf_buffer = create_student_report_pdf_reportlab(
                                    student_name,
                                    submission.get("user_email", ""),
                                    student_class,
                                    submission,
                                    questions,
                                    max_possible
                                )
                            except Exception as pdf_error:
                                # Fallback to FPDF nếu ReportLab fail
                                pdf_buffer = create_student_report_pdf_fpdf(
                                    student_name,
                                    submission.get("user_email", ""),
                                    student_class,
                                    submission,
                                    questions,
                                    max_possible
                                )
                            
                            if pdf_buffer is not None:
                                get_download_link_pdf(
                                    pdf_buffer, 
                                    f"bao_cao_{student_name.replace(' ', '_')}_{submission.get('id', '')}.pdf", 
                                    "📥 Tải xuống báo cáo chi tiết (PDF)"
                                )
                        except Exception as e:
                            st.error(f"Không thể tạo báo cáo PDF: {str(e)}")
    else:
        st.info("Không có dữ liệu học viên để hiển thị.")

def display_question_tab(submissions=None, questions=None):
    """Hiển thị tab phân tích câu hỏi"""
    if submissions is None:
        submissions = []
    if questions is None:
        questions = []
        
    st.subheader("Phân tích theo câu hỏi")
    
    # Thống kê tỷ lệ đúng/sai cho từng câu hỏi
    question_stats = {}
    
    for q in questions:
        q_id = str(q.get("id", ""))
        correct_count = 0
        wrong_count = 0
        skip_count = 0
        
        for s in submissions:
            # Đảm bảo responses đúng định dạng
            responses = s.get("responses", {})
            if isinstance(responses, str):
                try:
                    responses = json.loads(responses)
                except:
                    responses = {}
            
            user_ans = responses.get(q_id, [])
            
            # Đảm bảo user_ans là list
            if not isinstance(user_ans, list):
                if user_ans is not None:
                    user_ans = [user_ans]
                else:
                    user_ans = []
            
            if not user_ans:
                skip_count += 1
            else:
                # Sử dụng hàm từ database_helper (không dùng mock)
                try:
                    from database_helper import check_answer_correctness as db_check_answer
                    is_correct = db_check_answer(user_ans, q)
                except ImportError:
                    is_correct = check_answer_correctness(user_ans, q)
                
                if is_correct:
                correct_count += 1
            else:
                wrong_count += 1
        
        question_stats[q_id] = {
            "question": q.get("question", ""),
            "type": q.get("type", ""),  # Thêm thông tin loại câu hỏi
            "correct": correct_count,
            "wrong": wrong_count,
            "skip": skip_count,
            "total": correct_count + wrong_count + skip_count,
            "correct_rate": correct_count / (correct_count + wrong_count + skip_count) if (correct_count + wrong_count + skip_count) > 0 else 0
        }
    
    # DataFrame thống kê câu hỏi
    df_questions_data = [
        {
            "Câu hỏi ID": q_id,
            "Nội dung": stats["question"],
            "Loại câu hỏi": stats["type"],  # Thêm cột loại câu hỏi
            "Số lượng đúng": stats["correct"],
            "Số lượng sai": stats["wrong"],
            "Bỏ qua": stats["skip"],
            "Tổng số làm": stats["total"],
            "Tỷ lệ đúng (%)": f"{stats['correct_rate']*100:.1f}%"
        }
        for q_id, stats in question_stats.items()
    ]
    
    if not df_questions_data:
        st.info("Không có dữ liệu câu hỏi để phân tích.")
        return pd.DataFrame()
    
    df_questions = pd.DataFrame(df_questions_data)
    
    # Tạo bộ lọc loại câu hỏi
    question_types = ["Tất cả", "Checkbox", "Combobox", "Essay"]
    selected_type = st.selectbox("Lọc theo loại câu hỏi:", question_types, key="filter_question_type_tab3")
    
    # Áp dụng bộ lọc
    filtered_df = df_questions
    if selected_type != "Tất cả":
        filtered_df = df_questions[df_questions["Loại câu hỏi"] == selected_type]
    
    # Vẽ biểu đồ tỷ lệ đúng theo từng câu hỏi (chỉ cho các câu hỏi không phải tự luận)
    non_essay_df = filtered_df[filtered_df["Loại câu hỏi"] != "Essay"] if selected_type == "Tất cả" else filtered_df
    
    if not non_essay_df.empty:
        # Tạo dữ liệu cho biểu đồ
        q_ids = non_essay_df["Câu hỏi ID"].tolist()
        correct_rates = [float(rate.strip('%')) for rate in non_essay_df["Tỷ lệ đúng (%)"].tolist()]
        
        # Giới hạn độ dài câu hỏi để hiển thị trên biểu đồ
        short_questions = [f"Câu {q_id}: {non_essay_df[non_essay_df['Câu hỏi ID'] == q_id]['Nội dung'].values[0][:30]}..." for q_id in q_ids]
        
        # Tạo biểu đồ với kích thước nhỏ hơn
        fig, ax = plt.subplots(figsize=(10, 6))
        bars = ax.bar(short_questions, correct_rates, color='skyblue')
        
        # Xoay nhãn để tránh chồng chéo
        plt.xticks(rotation=45, ha='right')
        
        # Thêm nhãn giá trị
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                    f'{height:.1f}%', ha='center', va='bottom', fontsize=9)
        
        ax.set_ylim(0, 105)  # Giới hạn trục y từ 0-100%
        ax.set_xlabel("Câu hỏi")
        ax.set_ylabel("Tỷ lệ đúng (%)")
        ax.set_title("Tỷ lệ trả lời đúng theo từng câu hỏi")
        ax.grid(axis='y', linestyle='--', alpha=0.7)
        fig.set_constrained_layout(True)
        st.pyplot(fig)
    
    # Hiển thị biểu đồ riêng cho câu hỏi Essay nếu đã lọc theo Essay
    if selected_type == "Essay":
        st.info("Câu hỏi tự luận được đánh giá dựa trên việc học viên có trả lời hay không.")
        
        # Tạo dữ liệu cho biểu đồ tỷ lệ trả lời câu hỏi tự luận
        essay_df = filtered_df.copy()
        essay_q_ids = essay_df["Câu hỏi ID"].tolist()
        
        if essay_q_ids:
            # Tạo dữ liệu cho biểu đồ tròn tỷ lệ trả lời
            answer_rates = []
            question_texts = []
            
            for q_id in essay_q_ids:
                q_data = essay_df[essay_df["Câu hỏi ID"] == q_id].iloc[0]
                total = q_data["Tổng số làm"]
                answered = q_data["Số lượng đúng"]  # Đối với essay, "đúng" nghĩa là "đã trả lời"
                
                if total > 0:
                    answer_rate = (answered / total) * 100
                else:
                    answer_rate = 0
                
                answer_rates.append(answer_rate)
                question_texts.append(f"Câu {q_id}: {q_data['Nội dung'][:30]}...")
            
            # Vẽ biểu đồ tỷ lệ trả lời cho câu hỏi tự luận
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = ax.bar(question_texts, answer_rates, color='lightgreen')
            
            # Xoay nhãn để tránh chồng chéo
            plt.xticks(rotation=45, ha='right')
            
            # Thêm nhãn giá trị
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                        f'{height:.1f}%', ha='center', va='bottom', fontsize=9)
            
            ax.set_ylim(0, 105)  # Giới hạn trục y từ 0-100%
            ax.set_xlabel("Câu hỏi tự luận")
            ax.set_ylabel("Tỷ lệ trả lời (%)")
            ax.set_title("Tỷ lệ trả lời cho các câu hỏi tự luận")
            ax.grid(axis='y', linestyle='--', alpha=0.7)
            fig.set_constrained_layout(True)
            st.pyplot(fig)
    
    # Hiển thị bảng thống kê
    st.dataframe(filtered_df, width='stretch', hide_index=True)
    
    # Chi tiết từng câu hỏi
    if not filtered_df.empty:
        selected_question_options = [(f"Câu {row['Câu hỏi ID']}: {row['Nội dung']}") for _, row in filtered_df.iterrows()]
        if selected_question_options:
            selected_question = st.selectbox(
                "Chọn câu hỏi để xem chi tiết:",
                options=selected_question_options,
                key="question_select_tab3"
            )
            
            if selected_question:
                try:
                    q_id = selected_question.split(":")[0].replace("Câu ", "").strip()
                    q_data = question_stats[q_id]
                    q_detail = next((q for q in questions if str(q.get("id", "")) == q_id), None)
                    
                    if q_detail:
                        st.write(f"**{selected_question}**")
                        st.write(f"**Loại câu hỏi:** {q_data['type']}")
                        
                        # Hiển thị thống kê
                        col1, col2, col3, col4 = st.columns(4)
                        
                        if q_data['type'] == "Essay":
                            col1.metric("✍️ Đã trả lời", q_data["correct"])
                            col2.metric("⏭️ Không trả lời", q_data["skip"] + q_data["wrong"])
                        else:
                            col1.metric("✅ Đúng", q_data["correct"])
                            col2.metric("❌ Sai", q_data["wrong"])
                            col3.metric("⏭️ Bỏ qua", q_data["skip"])
                        
                        col4.metric("📊 Tỷ lệ đúng", f"{q_data['correct_rate']*100:.1f}%")
                        
                        # Tạo biểu đồ tròn cho phân phối đáp án
                        fig, ax = plt.subplots(figsize=(6, 4))
                        
                        if q_data['type'] == "Essay":
                            # Đối với câu hỏi tự luận, chỉ có hai loại: Đã trả lời và Không trả lời
                            labels = ['Đã trả lời', 'Không trả lời']
                            sizes = [q_data["correct"], q_data["skip"] + q_data["wrong"]]
                            colors = ['#4CAF50', '#9E9E9E']
                        else:
                            # Đối với câu hỏi trắc nghiệm
                            labels = ['Đúng', 'Sai', 'Bỏ qua']
                            sizes = [q_data["correct"], q_data["wrong"], q_data["skip"]]
                            colors = ['#4CAF50', '#F44336', '#9E9E9E']
                        
                        # Chỉ hiển thị phần trăm nếu giá trị > 0
                        patches, texts, autotexts = ax.pie(
                            sizes, 
                            labels=None,  # Không hiển thị nhãn trên biểu đồ
                            colors=colors, 
                            autopct=lambda p: f'{p:.1f}%' if p > 0 else '',
                            startangle=90,
                            pctdistance=0.85  # Đặt phần trăm gần hơn với trung tâm
                        )
                        
                        # Thiết lập kích thước font nhỏ hơn
                        for autotext in autotexts:
                            autotext.set_fontsize(9)
                        
                        # Thêm chú thích bên ngoài biểu đồ
                        ax.legend(labels, loc="upper right", fontsize=9)
                        
                        # Vẽ vòng tròn trắng ở giữa
                        centre_circle = plt.Circle((0, 0), 0.5, fc='white')
                        ax.add_patch(centre_circle)
                        
                        ax.axis('equal')  # Giữ tỷ lệ vòng tròn
                        fig.set_constrained_layout(True)
                        st.pyplot(fig)
                        
                        # Hiển thị thông tin thêm theo loại câu hỏi
                        if q_data['type'] == "Essay":
                            # Đối với câu hỏi tự luận
                            if q_detail.get("answer_template"):
                                st.write("**Mẫu câu trả lời:**")
                                st.text_area("", value=q_detail.get("answer_template", ""), 
                                            height=150, disabled=True, key=f"view_template_{q_id}")
                            else:
                                st.write("**Mẫu câu trả lời:** Không có")
                        else:
                            # Đối với câu hỏi trắc nghiệm, hiển thị đáp án đúng
                            st.write("**Đáp án đúng:**")
                            
                            # Chuẩn bị dữ liệu đáp án đúng
                            q_correct = q_detail.get("correct", [])
                            q_answers = q_detail.get("answers", [])
                            
                            if isinstance(q_correct, str):
                                try:
                                    q_correct = json.loads(q_correct)
                                except:
                                    try:
                                        q_correct = [int(x.strip()) for x in q_correct.split(",")]
                                    except:
                                        q_correct = []
                            
                            if isinstance(q_answers, str):
                                try:
                                    q_answers = json.loads(q_answers)
                                except:
                                    q_answers = [q_answers]
                            
                            try:
                                for i in q_correct:
                                    st.write(f"- {q_answers[i-1]}")
                            except (IndexError, TypeError):
                                st.write("- Lỗi hiển thị đáp án")
                except Exception as e:
                    st.error(f"Lỗi khi hiển thị chi tiết câu hỏi: {str(e)}")
    
    return df_questions

def display_student_list_tab(submissions=None, students=None, max_possible=0):
    """Hiển thị tab danh sách học viên"""
    if submissions is None:
        submissions = []
    if students is None:
        students = []
        
    st.subheader("Danh sách học viên")
    
    # Đảm bảo load lại students nếu chưa có - bao gồm tất cả roles
    if not students:
        try:
            # Thử dùng hàm get_all_students nếu có
            try:
                from database_helper import get_all_students
                students = get_all_students()
            except ImportError:
                # Fallback: load tất cả users với các role
                students = get_all_users(role=["Học viên", "student", "admin"])
            if not students:
                # Fallback cuối: load tất cả users và filter
                all_users = get_all_users(role=None)
                if all_users:
                    valid_roles = ["Học viên", "student", "admin"]
                    students = [u for u in all_users if u.get("role") in valid_roles]
                
                if not students:
                    st.warning("⚠️ Không thể load danh sách users từ Supabase. Vui lòng kiểm tra kết nối.")
                    st.info("💡 Gợi ý: Kiểm tra xem bảng 'users' trong Supabase có dữ liệu không.")
                    return pd.DataFrame(), pd.DataFrame()
        except Exception as e:
            st.error(f"❌ Lỗi khi load danh sách users: {str(e)}")
            import traceback
            st.code(traceback.format_exc(), language="python")
        return pd.DataFrame(), pd.DataFrame()
    
    # Phân tích roles
    role_counts = {}
    for s in students:
        role = s.get("role", "Unknown")
        role_counts[role] = role_counts.get(role, 0) + 1
    role_info = ", ".join([f"{r}: {c}" for r, c in role_counts.items()])
    st.info(f"📋 Tổng số users: {len(students)} ({role_info})")
    
    # Chuẩn bị dữ liệu - Đảm bảo HIỂN THỊ TẤT CẢ học viên (kể cả chưa làm bài)
    student_data = []
    for student in students:
        try:
            # Tìm tất cả bài nộp của học viên
            student_email = student.get("email", "")
            if not student_email:
                continue  # Bỏ qua nếu không có email
                
            student_submissions = [s for s in submissions if s.get("user_email") == student_email]
            submission_count = len(student_submissions)
            
            # Tìm điểm cao nhất
            max_student_score = max([s.get("score", 0) for s in student_submissions]) if student_submissions else 0
            
            # Thời gian đăng ký
            registration_date = format_date(student.get("registration_date"))
            
            # Đảm bảo lấy đầy đủ thông tin từ student dict
            full_name = student.get("full_name", "") or "Chưa có tên"
            class_name = student.get("class", "") or "Chưa phân lớp"
            
            student_data.append({
                "full_name": full_name,
                "email": student_email,
                "class": class_name,
                "registration_date": registration_date,
                "submission_count": submission_count,
                "max_score": max_student_score,
                "max_possible": max_possible,
                "percent": f"{(max_student_score/max_possible*100):.1f}%" if max_possible > 0 and max_student_score > 0 else ("0%" if max_possible > 0 else "N/A")
            })
        except Exception as e:
            st.warning(f"⚠️ Lỗi khi xử lý dữ liệu học viên {student.get('email', 'N/A')}: {str(e)}")
            # Vẫn thêm vào danh sách với dữ liệu cơ bản
            try:
                student_data.append({
                    "full_name": student.get("full_name", "Lỗi"),
                    "email": student.get("email", "N/A"),
                    "class": student.get("class", "Lỗi"),
                    "registration_date": format_date(student.get("registration_date")),
                    "submission_count": 0,
                    "max_score": 0,
                    "max_possible": max_possible,
                    "percent": "N/A"
                })
            except:
                pass
    
    # Đảm bảo có dữ liệu để hiển thị
    if not student_data:
        st.warning("⚠️ Không có dữ liệu học viên để hiển thị.")
        # Vẫn tạo DataFrame trống với cột đầy đủ
        df_students_list = pd.DataFrame(columns=["Họ và tên", "Email", "Lớp", "Ngày đăng ký", "Số lần làm bài", "Điểm cao nhất", "Điểm tối đa", "Tỷ lệ đúng"])
        df_class_stats = pd.DataFrame()
        return df_students_list, df_class_stats
    
    # DataFrame cho danh sách học viên
    students_list_data = [
        {
            "Họ và tên": s.get("full_name", "Chưa có tên"),
            "Email": s.get("email", "N/A"),
            "Lớp": s.get("class", "Chưa phân lớp"),
            "Ngày đăng ký": s.get("registration_date", "N/A"),
            "Số lần làm bài": s.get("submission_count", 0),
            "Điểm cao nhất": s.get("max_score", 0),
            "Điểm tối đa": s.get("max_possible", 0),
            "Tỷ lệ đúng": s.get("percent", "N/A")
        } for s in student_data
    ]
    
    df_students_list = pd.DataFrame(students_list_data)
    
    # Lọc theo lớp
    class_options = ["Tất cả"] + sorted(list(set([s["class"] for s in student_data if s["class"]])))
    class_filter = st.selectbox(
        "Lọc theo lớp:",
        options=class_options,
        key="class_filter_tab4"
    )
    
    df_students = pd.DataFrame(student_data)
    
    if class_filter != "Tất cả":
        df_students = df_students[df_students["class"] == class_filter]
    
    # Sắp xếp theo tên
    df_students = df_students.sort_values(by="full_name")
    
    # Hiển thị bảng
    st.dataframe(
        df_students,
        width='stretch',
        hide_index=True
    )
    
    # Thống kê theo lớp
    st.subheader("Thống kê theo lớp")
    
    # Nhóm theo lớp
    df_class_stats = pd.DataFrame()
    if not df_students.empty and "class" in df_students.columns:
        # Đảm bảo rằng class không rỗng
        df_students["class"] = df_students["class"].fillna("Không xác định")
        
        class_stats = df_students.groupby("class").agg({
            "email": "count",
            "submission_count": "sum",
            "max_score": "mean"
        }).reset_index()
        
        class_stats.columns = ["Lớp", "Số học viên", "Tổng số bài nộp", "Điểm trung bình"]
        class_stats["Điểm trung bình"] = class_stats["Điểm trung bình"].round(2)
        
        # DataFrame thống kê lớp
        df_class_stats = class_stats.copy()
        
        st.dataframe(
            class_stats,
            width='stretch',
            hide_index=True
        )
        
        # Biểu đồ cột nhỏ hơn cho số học viên theo lớp
        fig, ax = plt.subplots(figsize=(10, 4))
        ax.bar(class_stats["Lớp"], class_stats["Số học viên"], color='skyblue')
        ax.set_xlabel("Lớp")
        ax.set_ylabel("Số học viên")
        ax.set_title("Số học viên theo lớp")
        plt.xticks(rotation=45, ha='right')
        fig.set_constrained_layout(True)
        st.pyplot(fig)
    else:
        st.info("Không có đủ dữ liệu để hiển thị thống kê theo lớp.")
    
    return df_students_list, df_class_stats

def display_export_tab(df_all_submissions=None, df_questions=None, df_students_list=None, df_class_stats=None):
    """Hiển thị tab xuất báo cáo"""
    if df_all_submissions is None:
        df_all_submissions = pd.DataFrame()
    if df_questions is None:
        df_questions = pd.DataFrame()
    if df_students_list is None:
        df_students_list = pd.DataFrame()
    if df_class_stats is None:
        df_class_stats = pd.DataFrame()
        
    st.subheader("Xuất báo cáo")
    
    # Thêm tab cho các loại báo cáo khác nhau
    report_tab1, report_tab2 = st.tabs(["Báo cáo tổng hợp", "Báo cáo theo học viên"])
    
    with report_tab1:
        # Hiển thị các loại báo cáo có thể xuất
        if not df_all_submissions.empty:
            st.write("### 1. Báo cáo tất cả bài nộp")
            
            col1, col2 = st.columns(2)
            
            with col1:
                try:
                    # DOCX
                    docx_buffer = dataframe_to_docx(df_all_submissions, "Báo cáo tất cả bài nộp", "bao_cao_tat_ca_bai_nop.docx")
                    if docx_buffer is not None:
                        get_download_link_docx(docx_buffer, "bao_cao_tat_ca_bai_nop.docx", "📥 Tải xuống báo cáo (DOCX)")
                except Exception as e:
                    st.error(f"Lỗi khi tạo DOCX: {str(e)}")
            
            with col2:
                try:
                    # PDF - sử dụng FPDF thay vì ReportLab
                    try:
                        pdf_buffer = dataframe_to_pdf_reportlab(df_all_submissions, "Báo cáo tất cả bài nộp", "bao_cao_tat_ca_bai_nop.pdf")
                    except:
                    pdf_buffer = dataframe_to_pdf_fpdf(df_all_submissions, "Báo cáo tất cả bài nộp", "bao_cao_tat_ca_bai_nop.pdf")
                    if pdf_buffer is not None:
                        get_download_link_pdf(pdf_buffer, "bao_cao_tat_ca_bai_nop.pdf", "📥 Tải xuống báo cáo (PDF)")
                except Exception as e:
                    st.error(f"Lỗi khi tạo PDF: {str(e)}")
        
        # ... [tiếp tục cho các phần báo cáo khác]
        if not df_questions.empty:
            st.write("### 2. Báo cáo thống kê câu hỏi")
            
            col1, col2 = st.columns(2)
            
            with col1:
                try:
                    # DOCX
                    docx_buffer = dataframe_to_docx(df_questions, "Báo cáo thống kê câu hỏi", "bao_cao_thong_ke_cau_hoi.docx")
                    if docx_buffer is not None:
                        get_download_link_docx(docx_buffer, "bao_cao_thong_ke_cau_hoi.docx", "📥 Tải xuống báo cáo (DOCX)")
                except Exception as e:
                    st.error(f"Lỗi khi tạo DOCX: {str(e)}")
            
            with col2:
                try:
                    # PDF
                    try:
                        pdf_buffer = dataframe_to_pdf_reportlab(df_questions, "Báo cáo thống kê câu hỏi", "bao_cao_thong_ke_cau_hoi.pdf")
                    except:
                    pdf_buffer = dataframe_to_pdf_fpdf(df_questions, "Báo cáo thống kê câu hỏi", "bao_cao_thong_ke_cau_hoi.pdf")
                    if pdf_buffer is not None:
                        get_download_link_pdf(pdf_buffer, "bao_cao_thong_ke_cau_hoi.pdf", "📥 Tải xuống báo cáo (PDF)")
                except Exception as e:
                    st.error(f"Lỗi khi tạo PDF: {str(e)}")
        
        if not df_students_list.empty:
            st.write("### 3. Báo cáo danh sách học viên")
            
            col1, col2 = st.columns(2)
            
            with col1:
                try:
                    # DOCX
                    docx_buffer = dataframe_to_docx(df_students_list, "Báo cáo danh sách học viên", "bao_cao_danh_sach_hoc_vien.docx")
                    if docx_buffer is not None:
                        get_download_link_docx(docx_buffer, "bao_cao_danh_sach_hoc_vien.docx", "📥 Tải xuống báo cáo (DOCX)")
                except Exception as e:
                    st.error(f"Lỗi khi tạo DOCX: {str(e)}")
            
            with col2:
                try:
                    # PDF
                    try:
                        pdf_buffer = dataframe_to_pdf_reportlab(df_students_list, "Báo cáo danh sách học viên", "bao_cao_danh_sach_hoc_vien.pdf")
                    except:
                    pdf_buffer = dataframe_to_pdf_fpdf(df_students_list, "Báo cáo danh sách học viên", "bao_cao_danh_sach_hoc_vien.pdf")
                    if pdf_buffer is not None:
                        get_download_link_pdf(pdf_buffer, "bao_cao_danh_sach_hoc_vien.pdf", "📥 Tải xuống báo cáo (PDF)")
                except Exception as e:
                    st.error(f"Lỗi khi tạo PDF: {str(e)}")
        
        if not df_class_stats.empty:
            st.write("### 4. Báo cáo thống kê theo lớp")
            
            col1, col2 = st.columns(2)
            
            with col1:
                try:
                    # DOCX
                    docx_buffer = dataframe_to_docx(df_class_stats, "Báo cáo thống kê theo lớp", "bao_cao_thong_ke_lop.docx")
                    if docx_buffer is not None:
                        get_download_link_docx(docx_buffer, "bao_cao_thong_ke_lop.docx", "📥 Tải xuống báo cáo (DOCX)")
                except Exception as e:
                    st.error(f"Lỗi khi tạo DOCX: {str(e)}")
            
            with col2:
                try:
                    # PDF
                    try:
                        pdf_buffer = dataframe_to_pdf_reportlab(df_class_stats, "Báo cáo thống kê theo lớp", "bao_cao_thong_ke_lop.pdf")
                    except:
                    pdf_buffer = dataframe_to_pdf_fpdf(df_class_stats, "Báo cáo thống kê theo lớp", "bao_cao_thong_ke_lop.pdf")
                    if pdf_buffer is not None:
                        get_download_link_pdf(pdf_buffer, "bao_cao_thong_ke_lop.pdf", "📥 Tải xuống báo cáo (PDF)")
                except Exception as e:
                    st.error(f"Lỗi khi tạo PDF: {str(e)}")
        
        st.write("### 5. Báo cáo tổng hợp (Excel)")
        
        try:
            # Chuẩn bị danh sách DataFrame và tên sheet
            dfs = []
            sheet_names = []
            
            if not df_all_submissions.empty:
                dfs.append(df_all_submissions)
                sheet_names.append("Tất cả bài nộp")
            
            if not df_questions.empty:
                dfs.append(df_questions)
                sheet_names.append("Thống kê câu hỏi")
            
            if not df_students_list.empty:
                dfs.append(df_students_list)
                sheet_names.append("Danh sách học viên")
            
            if not df_class_stats.empty:
                dfs.append(df_class_stats)
                sheet_names.append("Thống kê lớp")
            
            if dfs and sheet_names:
                # Lấy questions và submissions để tính tổng hợp
                try:
                    from database_helper import get_all_questions, get_all_submissions
                    questions_data = get_all_questions()
                    submissions_data = get_all_submissions()
                except:
                    questions_data = None
                    submissions_data = None
                
                # Gọi export_to_excel - giờ nó tự tạo download button với tổng hợp
                export_to_excel(dfs, sheet_names, "bao_cao_tong_hop.xlsx", 
                               include_summary=True, 
                               questions=questions_data, 
                               submissions=submissions_data)
            else:
                st.info("Không có đủ dữ liệu để tạo báo cáo Excel.")
            
        except Exception as e:
            st.error(f"Lỗi khi tạo file Excel: {str(e)}")

    
    with report_tab2:
        st.write("### Báo cáo chi tiết theo từng học viên")
        
        # Lấy danh sách học viên và bài nộp từ database
        try:
            # Lấy dữ liệu từ database
            supabase = get_supabase_client()
            if not supabase:
                st.error("Không thể kết nối đến Supabase.")
                return
                
            # Load tất cả users (Học viên, student, admin)
            try:
                from database_helper import get_all_students
                students = get_all_students()
            except ImportError:
                students = get_all_users(role=["Học viên", "student", "admin"])
            questions = get_all_questions()
            max_possible = sum([q.get("score", 0) for q in questions])
            
            # Lấy danh sách email học viên từ các bài nộp
            student_emails = []
            for student in students:
                student_email = student.get("email", "")
                if student_email:
                    student_emails.append(student_email)
            
            student_emails = sorted(student_emails)
                
            if not student_emails:
                st.info("Không có dữ liệu học viên để hiển thị.")
                return
            
            # Hiển thị dropdown để chọn học viên
            selected_email = st.selectbox(
                "Chọn email học viên:",
                options=student_emails
            )
            
            if selected_email:
                # Lấy thông tin học viên
                student_info = next((student for student in students if student.get("email") == selected_email), None)
                if not student_info:
                    st.warning(f"Không tìm thấy thông tin học viên: {selected_email}")
                    return
                    
                student_name = student_info.get("full_name", "Không xác định")
                student_class = student_info.get("class", "Không xác định")
                
                # Lấy tất cả bài nộp của học viên này
                student_submissions = get_user_submissions(selected_email)
                
                if student_submissions:
                    st.success(f"Đã tìm thấy {len(student_submissions)} bài làm của học viên {student_name} ({selected_email})")
                    
                    # Hiển thị thông tin tổng quan
                    max_score = max([s.get("score", 0) for s in student_submissions]) if student_submissions else 0
                    best_submission = max(student_submissions, key=lambda x: x.get("score", 0))
                    best_score = best_submission.get("score", 0)
                    best_percent = (best_score / max_possible) * 100 if max_possible > 0 else 0
                    
                    col1, col2 = st.columns(2)
                    col1.metric("Số lần làm bài", len(student_submissions))
                    col2.metric("Điểm cao nhất", f"{best_score}/{max_possible} ({best_percent:.1f}%)")
                    
                    # Tạo DataFrame cho xuất báo cáo
                    student_report_data = []
                    
                    for idx, submission in enumerate(student_submissions):
                        # Xử lý timestamp
                        submission_time = "Không xác định"
                        if isinstance(submission.get("timestamp"), (int, float)):
                            try:
                                submission_time = datetime.fromtimestamp(submission.get("timestamp")).strftime("%H:%M:%S %d/%m/%Y")
                            except:
                                pass
                        else:
                            try:
                                dt = datetime.fromisoformat(submission.get("timestamp", "").replace("Z", "+00:00"))
                                submission_time = dt.strftime("%H:%M:%S %d/%m/%Y")
                            except:
                                pass
                        
                        # Đảm bảo responses đúng định dạng - parse từ JSON string từ database
                        responses = submission.get("responses", {})
                        if isinstance(responses, str):
                            try:
                                responses = json.loads(responses)
                            except json.JSONDecodeError as e:
                                print(f"Lỗi khi parse responses JSON trong display_export_tab: {e}")
                                responses = {}
                        
                        # Đảm bảo responses là dict
                        if not isinstance(responses, dict):
                            responses = {}
                        
                        # Tính số câu trả lời đúng - sử dụng dữ liệu thực từ database
                        correct_count = 0
                        for q in questions:
                            q_id = str(q.get("id", ""))
                            user_ans = responses.get(q_id, [])
                            
                            # Đảm bảo user_ans là list
                            if not isinstance(user_ans, list):
                                if user_ans is not None:
                                    user_ans = [user_ans]
                                else:
                                    user_ans = []
                            
                            # Sử dụng hàm từ database_helper (không dùng mock)
                            try:
                                from database_helper import check_answer_correctness as db_check_answer
                                is_correct = db_check_answer(user_ans, q)
                            except ImportError:
                            is_correct = check_answer_correctness(user_ans, q)
                            
                            if is_correct:
                                correct_count += 1
                        
                        score_percent = (submission.get("score", 0) / max_possible) * 100 if max_possible > 0 else 0
                        
                        # Thêm dữ liệu vào danh sách
                        entry = {
                            "Lần làm": idx + 1,
                            "Thời gian": submission_time,
                            "Điểm số": submission.get("score", 0),
                            "Điểm tối đa": max_possible,
                            "Tỷ lệ đúng": f"{score_percent:.1f}%",
                            "Số câu đúng": f"{correct_count}/{len(questions)}"
                        }
                        
                        # Thêm chi tiết từng câu hỏi
                        for q in questions:
                            q_id = str(q.get("id", ""))
                            user_ans = responses.get(q_id, [])
                            is_correct = check_answer_correctness(user_ans, q)
                            
                            entry[f"Câu {q_id}"] = ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời"
                            entry[f"Câu {q_id} - Kết quả"] = "Đúng" if is_correct else "Sai"
                        
                        student_report_data.append(entry)
                    
                    # Tạo DataFrame cho báo cáo
                    df_student_report = pd.DataFrame(student_report_data)
                    
                    # Hiển thị dữ liệu dạng bảng
                    st.write("### Chi tiết các lần làm bài")
                    st.dataframe(df_student_report, hide_index=True, width='stretch')
                    
                    # Tạo báo cáo Word cho học viên này
                    try:
                        # Tạo tiêu đề
                        title = f"Báo cáo chi tiết học viên: {student_name} ({selected_email})"
                        
                        st.write("### Tải xuống báo cáo học viên")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            # Word
                            try:
                            docx_buffer = dataframe_to_docx(df_student_report, title, f"bao_cao_{student_name}.docx")
                                if docx_buffer is None:
                                    st.error("Không thể tạo báo cáo DOCX: Buffer rỗng")
                                else:
                            st.markdown(
                                get_download_link_docx(docx_buffer, 
                                                    f"bao_cao_{student_name.replace(' ', '_')}.docx", 
                                                    "Tải xuống báo cáo DOCX"), 
                                unsafe_allow_html=True
                            )
                            except Exception as e:
                                st.error(f"Lỗi khi tạo báo cáo DOCX: {str(e)}")
                        
                        with col2:
                            # PDF
                            try:
                                try:
                                    pdf_buffer = dataframe_to_pdf_reportlab(df_student_report, title, f"bao_cao_{student_name}.pdf")
                                except:
                            pdf_buffer = dataframe_to_pdf_fpdf(df_student_report, title, f"bao_cao_{student_name}.pdf")
                                if pdf_buffer is not None:
                                    get_download_link_pdf(
                                        pdf_buffer, 
                                                    f"bao_cao_{student_name.replace(' ', '_')}.pdf", 
                                        "📥 Tải xuống báo cáo PDF"
                            )
                            except Exception as e:
                                st.error(f"Lỗi khi tạo báo cáo PDF: {str(e)}")
                        
                        # Tạo báo cáo chi tiết cho từng lần làm
                        st.write("### Tải báo cáo chi tiết từng lần làm")
                        
                        # Thêm Excel export cho tất cả các lần làm
                        st.write("#### Tải xuống Excel tổng hợp các lần làm")
                        try:
                            # Tạo danh sách DataFrames cho từng lần làm
                            excel_dfs = []
                            excel_sheet_names = []
                            
                        for idx, submission in enumerate(student_submissions):
                                # Tạo DataFrame cho lần làm này
                                submission_data = []
                                
                                # Xử lý timestamp
                                submission_time = "Không xác định"
                                if isinstance(submission.get("timestamp"), (int, float)):
                                    try:
                                        submission_time = datetime.fromtimestamp(submission.get("timestamp")).strftime("%H:%M:%S %d/%m/%Y")
                                    except:
                                        pass
                                else:
                                    try:
                                        dt = datetime.fromisoformat(submission.get("timestamp", "").replace("Z", "+00:00"))
                                        submission_time = dt.strftime("%H:%M:%S %d/%m/%Y")
                                    except:
                                        pass
                                
                                # Đảm bảo responses đúng định dạng
                                responses = submission.get("responses", {})
                                if isinstance(responses, str):
                                    try:
                                        responses = json.loads(responses)
                                    except:
                                        responses = {}
                                if not isinstance(responses, dict):
                                    responses = {}
                                
                                # Tính toán chi tiết
                                for q in questions:
                                    q_id = str(q.get("id", ""))
                                    user_ans = responses.get(q_id, [])
                                    if not isinstance(user_ans, list):
                                        user_ans = [user_ans] if user_ans is not None else []
                                    
                                    try:
                                        from database_helper import check_answer_correctness as db_check_answer
                                        is_correct = db_check_answer(user_ans, q)
                                    except ImportError:
                                        is_correct = check_answer_correctness(user_ans, q)
                                    
                                    row_data = {
                                        "Câu hỏi ID": q_id,
                                        "Nội dung câu hỏi": q.get("question", ""),
                                        "Loại câu hỏi": q.get("type", ""),
                                        "Đáp án của học viên": ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời",
                                        "Đáp án đúng": ", ".join([str(q.get("answers", [])[i-1]) if i <= len(q.get("answers", [])) else "" for i in q.get("correct", [])]) if q.get("type") != "Essay" else "Câu hỏi tự luận",
                                        "Kết quả": "Đúng" if is_correct else "Sai",
                                        "Điểm": q.get("score", 0) if is_correct else 0
                                    }
                                    submission_data.append(row_data)
                                
                                df_submission = pd.DataFrame(submission_data)
                                excel_dfs.append(df_submission)
                                excel_sheet_names.append(f"Lan {idx+1} - {submission_time}")
                            
                            if excel_dfs:
                                export_to_excel(excel_dfs, excel_sheet_names, f"bao_cao_chi_tiet_{student_name.replace(' ', '_')}_tat_ca_lan_lam.xlsx")
                        except Exception as e:
                            st.error(f"Lỗi khi tạo báo cáo Excel: {str(e)}")
                        
                        for idx, submission in enumerate(student_submissions):
                            st.write(f"#### Lần làm {idx+1}")
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                try:
                                    # Word
                                    docx_buffer = create_student_report_docx(
                                        student_name,
                                        selected_email,
                                        student_class,
                                        submission,
                                        questions,
                                        max_possible
                                    )
                                    
                                    if docx_buffer is not None:
                                        get_download_link_docx(
                                            docx_buffer, 
                                            f"bao_cao_chi_tiet_{student_name.replace(' ', '_')}_lan_{idx+1}.docx", 
                                            f"📥 Tải xuống báo cáo lần {idx+1} (DOCX)"
                                    )
                                except Exception as e:
                                    st.error(f"Lỗi khi tạo báo cáo DOCX lần {idx+1}: {str(e)}")
                            
                            with col2:
                                try:
                                    # PDF
                                    try:
                                        pdf_buffer = create_student_report_pdf_reportlab(
                                            student_name,
                                            selected_email,
                                            student_class,
                                            submission,
                                            questions,
                                            max_possible
                                        )
                                    except Exception as pdf_error:
                                        # Fallback to FPDF nếu ReportLab fail
                                        pdf_buffer = create_student_report_pdf_fpdf(
                                            student_name,
                                            selected_email,
                                            student_class,
                                            submission,
                                            questions,
                                            max_possible
                                        )
                                    
                                    if pdf_buffer is not None:
                                        get_download_link_pdf(
                                            pdf_buffer, 
                                            f"bao_cao_chi_tiet_{student_name.replace(' ', '_')}_lan_{idx+1}.pdf", 
                                            f"📥 Tải xuống báo cáo lần {idx+1} (PDF)"
                                    )
                                except Exception as e:
                                    st.error(f"Lỗi khi tạo báo cáo PDF lần {idx+1}: {str(e)}")
                            
                            with col3:
                                try:
                                    # Excel cho lần làm này
                                    submission_data = []
                                    
                                    # Xử lý responses
                                    responses = submission.get("responses", {})
                                    if isinstance(responses, str):
                                        try:
                                            responses = json.loads(responses)
                                        except:
                                            responses = {}
                                    if not isinstance(responses, dict):
                                        responses = {}
                                    
                                    for q in questions:
                                        q_id = str(q.get("id", ""))
                                        user_ans = responses.get(q_id, [])
                                        if not isinstance(user_ans, list):
                                            user_ans = [user_ans] if user_ans is not None else []
                                        
                                        try:
                                            from database_helper import check_answer_correctness as db_check_answer
                                            is_correct = db_check_answer(user_ans, q)
                                        except ImportError:
                                            is_correct = check_answer_correctness(user_ans, q)
                                        
                                        row_data = {
                                            "Câu hỏi ID": q_id,
                                            "Nội dung câu hỏi": q.get("question", ""),
                                            "Loại câu hỏi": q.get("type", ""),
                                            "Đáp án của học viên": ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời",
                                            "Đáp án đúng": ", ".join([str(q.get("answers", [])[i-1]) if i <= len(q.get("answers", [])) else "" for i in q.get("correct", [])]) if q.get("type") != "Essay" else "Câu hỏi tự luận",
                                            "Kết quả": "Đúng" if is_correct else "Sai",
                                            "Điểm": q.get("score", 0) if is_correct else 0
                                        }
                                        submission_data.append(row_data)
                                    
                                    df_submission = pd.DataFrame(submission_data)
                                    excel_dfs_single = [df_submission]
                                    excel_sheet_names_single = [f"Chi tiet lan {idx+1}"]
                                    export_to_excel(excel_dfs_single, excel_sheet_names_single, f"bao_cao_chi_tiet_{student_name.replace(' ', '_')}_lan_{idx+1}.xlsx")
                                except Exception as e:
                                    st.error(f"Lỗi khi tạo báo cáo Excel lần {idx+1}: {str(e)}")
                        
                    except Exception as e:
                        st.error(f"Lỗi khi tạo báo cáo: {str(e)}")
                
                else:
                    st.warning(f"Không tìm thấy bài nộp nào của học viên {student_name} ({selected_email})")
            else:
                st.info("Vui lòng chọn email học viên để xem và xuất báo cáo")
                
        except Exception as e:
            st.error(f"Lỗi khi xử lý báo cáo theo học viên: {str(e)}")

def view_statistics():
    """Hiển thị trang thống kê và báo cáo"""
    st.title("📊 Báo cáo & thống kê")
    
    # Khởi tạo biến trước
    questions = []
    students = []
    submissions = []
    max_possible = 0
    df_questions = pd.DataFrame()
    df_students_list = pd.DataFrame()
    df_class_stats = pd.DataFrame()
    df_all_submissions = pd.DataFrame()
    
    try:
        # Lấy dữ liệu THỰC từ database - KHÔNG dùng mock/fake data
        # Đảm bảo import từ database_helper (nơi có hàm thực)
        try:
            from database_helper import get_all_questions as db_get_all_questions
            questions = db_get_all_questions()
            print(f"✓ Đã load {len(questions)} câu hỏi từ database (database_helper)")
        except ImportError:
            print("⚠️ Warning: Không thể import get_all_questions từ database_helper, dùng fallback")
            questions = get_all_questions()  # Fallback to mock (chỉ khi không có database_helper)
        
        # Validate và normalize questions - đảm bảo parse JSON đúng cách
        validated_questions = []
        for q in questions:
            if not isinstance(q, dict):
                print(f"Warning: Bỏ qua câu hỏi không hợp lệ (không phải dict)")
                continue
            
            # Parse answers từ JSON string nếu cần (từ database)
            if isinstance(q.get("answers"), str):
                try:
                    q["answers"] = json.loads(q["answers"])
                except:
                    try:
                        q["answers"] = [q["answers"]]
                    except:
                        q["answers"] = []
            
            # Parse correct từ JSON string nếu cần (từ database)
            if isinstance(q.get("correct"), str):
                try:
                    q["correct"] = json.loads(q["correct"])
                except:
                    try:
                        q["correct"] = [int(x.strip()) for x in q["correct"].split(",")]
                    except:
                        q["correct"] = []
            
            # Đảm bảo là list
            if not isinstance(q.get("answers"), list):
                q["answers"] = []
            if not isinstance(q.get("correct"), list):
                q["correct"] = []
            
            validated_questions.append(q)
        questions = validated_questions
        
        if not questions:
            st.warning("⚠️ Không có câu hỏi nào được load từ database. Kiểm tra kết nối Supabase.")
        
        # Lấy TẤT CẢ users từ database (bao gồm "Học viên", "student", "admin")
        try:
            from database_helper import get_all_students
            students = get_all_students()
            print(f"✓ Đã load {len(students)} users từ database (get_all_students)")
        except ImportError:
            # Fallback nếu hàm chưa có
            students = get_all_users(role=["Học viên", "student", "admin"])
            print(f"✓ Đã load {len(students)} users từ database (get_all_users - fallback)")
        if not students:
            # Thử load từng role riêng
            students_hv = get_all_users(role="Học viên") if 'get_all_users' in globals() else []
            students_st = get_all_users(role="student") if 'get_all_users' in globals() else []
            students_ad = get_all_users(role="admin") if 'get_all_users' in globals() else []
            students = students_hv + students_st + students_ad
        
        # Debug: hiển thị số lượng users được load
        if students:
            role_counts = {}
            for s in students:
                role = s.get("role", "Unknown")
                role_counts[role] = role_counts.get(role, 0) + 1
            role_info = ", ".join([f"{r}: {c}" for r, c in role_counts.items()])
            st.sidebar.info(f"📊 Đã load {len(students)} users từ database ({role_info})")
        else:
            st.sidebar.warning("⚠️ Không tìm thấy users nào trong database")
        
        # Tạo form tìm kiếm email nếu muốn xem báo cáo theo học viên cụ thể
        with st.sidebar:
            st.subheader("Tìm kiếm học viên")
            search_email = st.text_input("Nhập email học viên:", key="search_email_stats")
            search_button = st.button("Tìm kiếm", key="search_button_stats")
        
        if search_button and search_email:
            submissions = get_user_submissions(search_email)
            if not submissions:
                st.warning(f"Không tìm thấy bài nộp của học viên: {search_email}")
                return
        else:
            # Lấy tất cả bài nộp trực tiếp từ database (hiệu quả hơn) - KHÔNG dùng mock
            try:
                # Đảm bảo import từ database_helper
                try:
                    from database_helper import get_all_submissions as db_get_all_submissions
                    submissions = db_get_all_submissions()
                    print(f"✓ Đã load {len(submissions)} bài nộp từ database (database_helper)")
                except ImportError:
                    print("⚠️ Warning: Không thể import get_all_submissions từ database_helper")
                    submissions = get_all_submissions()  # Fallback
                
                # Validate và parse responses từ JSON cho mỗi submission
                validated_submissions = []
                for s in submissions:
                    if not isinstance(s, dict):
                        continue
                    
                    # Parse responses từ JSON string nếu cần (từ database)
                    if isinstance(s.get("responses"), str):
                        try:
                            s["responses"] = json.loads(s["responses"])
                        except json.JSONDecodeError as e:
                            print(f"Warning: Lỗi parse responses cho submission {s.get('id', 'N/A')}: {e}")
                            s["responses"] = {}
                    
                    # Đảm bảo responses là dict
                    if not isinstance(s.get("responses"), dict):
                        s["responses"] = {}
                    
                    validated_submissions.append(s)
                
                submissions = validated_submissions
                
            except Exception as e:
                st.error(f"Lỗi khi lấy dữ liệu bài nộp từ Supabase: {str(e)}")
                import traceback
                traceback.print_exc()
                # Fallback: thử lấy từng học viên nếu lỗi
                st.info("Đang thử cách khác...")
                submissions = []
            for student in students:
                try:
                        try:
                            from database_helper import get_user_submissions as db_get_user_submissions
                            student_submissions = db_get_user_submissions(student.get("email", ""))
                        except ImportError:
                    student_submissions = get_user_submissions(student.get("email", ""))
                        
                        if student_submissions:
                            # Validate responses cho mỗi submission
                            for sub in student_submissions:
                                if isinstance(sub.get("responses"), str):
                                    try:
                                        sub["responses"] = json.loads(sub["responses"])
                                    except:
                                        sub["responses"] = {}
                    submissions.extend(student_submissions)
                    except Exception as ex:
                        print(f"Lỗi khi lấy dữ liệu của học viên {student.get('email', '')}: {str(ex)}")
        
        if not questions:
            st.warning("⚠️ Chưa có dữ liệu câu hỏi nào trong hệ thống. Vui lòng thêm câu hỏi trước.")
            return
        
        if not submissions:
            st.info("ℹ️ Chưa có ai nộp khảo sát. Các tab thống kê sẽ hiển thị khi có dữ liệu.")
            # Vẫn hiển thị tab nhưng với thông báo không có dữ liệu
            tab1, tab2, tab3, tab4, tab5 = st.tabs(["Tổng quan", "Theo học viên", "Theo câu hỏi", "Danh sách học viên", "Xuất báo cáo"])
            with tab1:
                st.info("Chưa có dữ liệu bài nộp để hiển thị.")
            with tab2:
                st.info("Chưa có dữ liệu bài nộp để hiển thị.")
            with tab3:
                st.info("Chưa có dữ liệu bài nộp để hiển thị.")
            with tab4:
                if students:
                    df_students = pd.DataFrame([
                        {
                            "Họ và tên": s.get("full_name", ""),
                            "Email": s.get("email", ""),
                            "Lớp": s.get("class", ""),
                            "Role": s.get("role", ""),
                            "Số lần làm bài": 0
                        }
                        for s in students
                    ])
                    st.dataframe(df_students, width='stretch')
                else:
                    st.info("Chưa có users nào trong hệ thống.")
            with tab5:
                st.info("Chưa có dữ liệu để xuất báo cáo.")
            return
        
        # Tạo tab thống kê
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["Tổng quan", "Theo học viên", "Theo câu hỏi", "Danh sách học viên", "Xuất báo cáo"])
        
        # Tính tổng điểm tối đa
        max_possible = sum([q.get("score", 0) for q in questions])
        
        # Chuẩn bị dữ liệu cho tất cả các submissions
        all_submission_data = []
        
        for s in submissions:
            try:
                # Tìm thông tin học viên
                student_info = next((student for student in students if student.get("email") == s.get("user_email")), None)
                full_name = student_info.get("full_name", "Không xác định") if student_info else "Không xác định"
                class_name = student_info.get("class", "Không xác định") if student_info else "Không xác định"
                
                # Chuyển đổi timestamp sang định dạng đọc được
                submission_time = "Không xác định"
                if isinstance(s.get("timestamp"), (int, float)):
                    try:
                        submission_time = datetime.fromtimestamp(s.get("timestamp")).strftime("%d/%m/%Y %H:%M:%S")
                    except:
                        pass
                else:
                    try:
                        dt = datetime.fromisoformat(s.get("timestamp", "").replace("Z", "+00:00"))
                        submission_time = dt.strftime("%d/%m/%Y %H:%M:%S")
                    except:
                        pass
                
                # Thêm thông tin cơ bản
                submission_data = {
                    "ID": s.get("id", ""),
                    "Email": s.get("user_email", ""),
                    "Họ và tên": full_name,
                    "Lớp": class_name,
                    "Thời gian nộp": submission_time,
                    "Điểm số": s.get("score", 0),
                    "Điểm tối đa": max_possible,
                    "Tỷ lệ đúng": f"{(s.get('score', 0)/max_possible*100):.1f}%" if max_possible > 0 else "N/A"
                }
                
                # Chuyển đổi responses từ JSON string thành dict từ database
                responses = s.get("responses", {})
                if isinstance(responses, str):
                    try:
                        responses = json.loads(responses)
                    except json.JSONDecodeError as e:
                        print(f"Lỗi parse responses cho submission {s.get('id', 'N/A')}: {e}")
                        responses = {}
                
                # Đảm bảo responses là dict
                if not isinstance(responses, dict):
                    responses = {}
                
                # Thêm câu trả lời của từng câu hỏi - sử dụng dữ liệu thực từ database
                for q in questions:
                    q_id = str(q.get("id", ""))
                    user_ans = responses.get(q_id, [])
                    
                    # Đảm bảo user_ans là list
                    if not isinstance(user_ans, list):
                        if user_ans is not None:
                            user_ans = [user_ans]
                        else:
                            user_ans = []
                    
                    # Đảm bảo q["correct"] và q["answers"] có định dạng đúng (đã được normalize ở trên)
                    q_correct = q.get("correct", [])
                    q_answers = q.get("answers", [])
                    
                    # Parse nếu chưa được normalize (backup)
                    if isinstance(q_correct, str):
                        try:
                            q_correct = json.loads(q_correct)
                        except:
                            try:
                                q_correct = [int(x.strip()) for x in q_correct.split(",")]
                            except:
                                q_correct = []
                    
                    if isinstance(q_answers, str):
                        try:
                            q_answers = json.loads(q_answers)
                        except:
                            q_answers = [q_answers]
                    
                    try:
                        expected = [q_answers[i - 1] for i in q_correct]
                    except (IndexError, TypeError):
                        expected = ["Lỗi đáp án"]
                        
                    # Sử dụng hàm từ database_helper (không dùng mock)
                    try:
                        from database_helper import check_answer_correctness as db_check_answer
                        is_correct = db_check_answer(user_ans, q)
                    except ImportError:
                    is_correct = check_answer_correctness(user_ans, q)
                    
                    # Thêm thông tin câu hỏi
                    submission_data[f"Câu {q_id}: {q.get('question', '')}"] = ", ".join([str(a) for a in user_ans]) if user_ans else "Không trả lời"
                    submission_data[f"Câu {q_id} - Đúng/Sai"] = "Đúng" if is_correct else "Sai"
                
                all_submission_data.append(submission_data)
            except Exception as e:
                st.error(f"Lỗi khi xử lý submission ID {s.get('id', '')}: {str(e)}")
        
        # DataFrame chứa tất cả bài nộp
        df_all_submissions = pd.DataFrame(all_submission_data) if all_submission_data else pd.DataFrame()
        
        with tab1:
            display_overview_tab(submissions, students, questions, max_possible)
        
        with tab2:
            display_student_tab(submissions, students, questions, max_possible)
        
        with tab3:
            df_questions = display_question_tab(submissions, questions)
        
        with tab4:
            df_students_list, df_class_stats = display_student_list_tab(submissions, students, max_possible)
        
        with tab5:
            display_export_tab(df_all_submissions, df_questions, df_students_list, df_class_stats)
    
    except Exception as e:
        st.error(f"Đã xảy ra lỗi không mong muốn: {str(e)}")
        traceback.print_exc()



# Chỉ chạy hàm main khi chạy file này trực tiếp
if __name__ == "__main__":
    st.set_page_config(
        page_title="Báo cáo & Thống kê",
        page_icon="📊",
        layout="wide",
    )
    view_statistics()
